from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(r"D:\Workspace\Codex\Crypto trading")
JSON_PATH = ROOT / "outputs" / "nxt_crypto_btc_sol_sui_6y_v30_close_25_correct" / "nxt_v30_close_25_correct_6y_results.json"
DOCX_PATH = ROOT / "NXT_Trading_System_v3.0_Corrected.docx"


def fmt_r(value: float) -> str:
    return f"{value:+.2f}R"


def fmt_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Muc"
    table.rows[0].cells[1].text = "Gia tri"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def main() -> None:
    data = json.loads(JSON_PATH.read_text(encoding="utf-8"))
    stats = data["stats"]
    summary = data["summary"]

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("NXT Trading System v3.0 - Corrected", 0)
    p = doc.add_paragraph()
    p.add_run("Ban theo doi chinh trong project. ").bold = True
    p.add_run("Khong dung lai ket qua +91R vi do la ban loi logic da bi loai bo.")

    add_kv_table(doc, [
        ("Version", "NXT v3.0 Corrected"),
        ("Ngay tao file", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Backtest period", f'{data["period"]["start"][:10]} to {data["period"]["end"][:10]}'),
        ("Symbols", ", ".join(data["symbols"])),
        ("Exit baseline", "Corrected full close / BE logic at 2.5 ATR"),
        ("Trades", str(stats["trades"])),
        ("Total R", fmt_r(stats["totalR"])),
        ("Win rate", fmt_pct(stats["winRate"])),
        ("Average R / trade", fmt_r(stats["avgR"])),
        ("Max drawdown", fmt_r(stats["maxDrawdownR"])),
    ])

    doc.add_heading("Rule Summary", level=1)
    for line in [
        "Entry/filter logic follows the corrected NXT v3 backtest family.",
        "Daily timeframe; signals are evaluated on closed daily candles and entered at the next daily open.",
        "No weekly regime filter.",
        "Long entry: SSL bullish crossover, price crosses above EMA20 within last 3 candles, Net Volume > 0, and distance to EMA50 <= 2 ATR.",
        "Short entry: SSL bearish crossover, price crosses below EMA20 within last 3 candles, Net Volume < 0, and distance to EMA50 <= 2 ATR.",
        "Stop loss is 1.5 x ATR(14) from entry.",
        "Corrected exit benchmark uses the latest corrected 6-year result set. Use the workbook/JSON in project root for audit.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Per Symbol Result", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Symbol", "Trades", "Win Rate", "Total R", "Avg R", "Worst R"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in summary:
        cells = table.add_row().cells
        cells[0].text = row["symbol"].replace("USDT", "")
        cells[1].text = str(row["trades"])
        cells[2].text = fmt_pct(row["winRate"])
        cells[3].text = fmt_r(row["totalR"])
        cells[4].text = fmt_r(row["avgR"])
        cells[5].text = fmt_r(row["worstR"])

    doc.add_heading("Root Files", level=1)
    for line in [
        "NXT_Trading_System_v3.0_Corrected.docx",
        "NXT_V30_Corrected_Backtest_6Y_BTC_SOL_SUI.xlsx",
        "NXT_V30_Corrected_Backtest_6Y_BTC_SOL_SUI.json",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
