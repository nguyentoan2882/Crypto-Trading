from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(r"D:\Workspace\Codex\Crypto trading")
SRC_JSON = ROOT / "outputs" / "nxt31_utc7_continuation_6y" / "nxt31_utc7_continuation_6y_results.json"
SRC_XLSX = ROOT / "outputs" / "nxt31_utc7_continuation_6y" / "NXT31_UTC7_With_Continuation_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST = ROOT / "latest"
OUT_JSON = LATEST / "NXT_Latest_NXT32_UTC7_RunnerA_Continuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.json"
OUT_XLSX = LATEST / "NXT_Latest_NXT32_UTC7_RunnerA_Continuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
OUT_DOCX = LATEST / "NXT_Latest_NXT32_System_And_Indicators.docx"
SUMMARY = LATEST / "NXT_Latest_Summary.md"


def build_doc(data):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title = doc.add_heading("NXT v3.2 + Continuation", 0)
    doc.add_paragraph("Latest selected system: UTC+7 daily candles, Runner A, continuation module, no risk-off overlay.")

    doc.add_heading("Backtest Snapshot", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    stats = data["stats"]
    rows = [
        ("System", "NXT v3.2"),
        ("Timezone", "Asia/Saigon UTC+7 daily candles"),
        ("Total trades", stats["trades"]),
        ("Win rate", f"{stats['winRate']:.2%}"),
        ("Total R", f"{stats['totalR']:.2f}R"),
        ("Average R", f"{stats['avgR']:.2f}R"),
        ("Max DD R", f"{stats['maxDrawdownR']:.2f}R"),
        ("Continuation trades", data["continuationStats"]["trades"]),
        ("Continuation R", f"{data['continuationStats']['totalR']:.2f}R"),
    ]
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)

    doc.add_heading("Indicators", level=1)
    for text in [
        "SSL Channel 10/10: primary trend flip and runner exit signal.",
        "EMA20 and EMA50: EMA20 cross/reclaim and EMA50 distance filter.",
        "ATR(14): stop, TP1 trigger and distance normalization.",
        "RSI(14): regime/noise filter for primary and continuation entries.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Primary Entry Rules", level=1)
    for text in [
        "LONG: SSL flips bullish, price crosses above EMA20 within 3 candles, distance to EMA50 <= 2 ATR, RSI14 > 50.",
        "SHORT: SSL flips bearish, price crosses below EMA20 within 3 candles, distance to EMA50 <= 2 ATR, RSI14 < 50.",
        "Entry is next UTC+7 daily open after signal candle close.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Continuation Module", level=1)
    for text in [
        "Continuation is only added when no primary trade on the same symbol is already open.",
        "SSL re-entry must occur within the last 3 candles.",
        "Distance to EMA50 must be <= 2.5 ATR.",
        "LONG continuation requires RSI14 > 55 and close above EMA20.",
        "SHORT continuation requires RSI14 < 45 and close below EMA20.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Exit Rules", level=1)
    for text in [
        "Runner A: close 50% at TP1 = 2.5 ATR.",
        "Move stop on remaining 50% to breakeven after TP1.",
        "Remaining runner exits on opposite SSL flip or breakeven stop.",
        "No risk-off overlay is applied in v3.2.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.save(OUT_DOCX)


def main():
    LATEST.mkdir(parents=True, exist_ok=True)
    data = json.loads(SRC_JSON.read_text(encoding="utf-8"))
    packaged = {
        "systemVersion": "NXT v3.2 + UTC+7 + Runner A + Continuation + No Risk-Off",
        "source": str(SRC_JSON),
        "stats": data["stats"],
        "continuationStats": data["continuationStats"],
        "trades": data["trades"],
        "datasets": data["datasets"],
        "assumptions": data["assumptions"],
    }
    OUT_JSON.write_text(json.dumps(packaged, indent=2), encoding="utf-8")
    shutil.copy2(SRC_XLSX, OUT_XLSX)
    build_doc(packaged)

    SUMMARY.write_text(
        "\n".join(
            [
                "# Latest NXT System",
                "",
                "System: NXT v3.2 + UTC+7 + Runner A + Continuation + No Risk-Off",
                "",
                f"Total R: {data['stats']['totalR']:.2f}R",
                f"Max DD R: {data['stats']['maxDrawdownR']:.2f}R",
                f"Trades: {data['stats']['trades']}",
                f"Win rate: {data['stats']['winRate']:.2%}",
                f"Continuation trades: {data['continuationStats']['trades']}",
                f"Continuation R: {data['continuationStats']['totalR']:.2f}R",
                "",
                "Notes: UTC+7 daily candles are used to match the TradingView setup on the local machine. Risk-off is not applied in v3.2.",
                "",
                f"Workbook: {OUT_XLSX.name}",
                f"JSON: {OUT_JSON.name}",
                f"System doc: {OUT_DOCX.name}",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
