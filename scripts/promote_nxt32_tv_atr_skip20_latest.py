from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

import backtest_nxt31_utc7_latest as base
import backtest_nxt32_native_1d_latest as native
import rebuild_nxt32_native_1d_tv_atr_latest as tv_atr
from test_nxt32_tv_atr_skip_20pct_reversal import backtest_symbol


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "nxt32_tv_atr_skip20_latest_promoted"
OUT_JSON = OUT_DIR / "nxt32_tv_atr_skip20_latest_promoted_results.json"
OUT_XLSX = OUT_DIR / "NXT32_TV_ATR_Skip20_Latest_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST = ROOT / "latest"
LATEST_JSON = LATEST / "NXT_Latest_NXT32_Native1D_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.json"
LATEST_XLSX = LATEST / "NXT_Latest_NXT32_Native1D_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST_DOCX = LATEST / "NXT_Latest_NXT32_System_And_Indicators.docx"
LATEST_SUMMARY = LATEST / "NXT_Latest_Summary.md"


RULES = [
    "Data: Binance native 1D candles for BTCUSDT, SOLUSDT, and SUIUSDT.",
    "ATR14: TradingView default ATR using Wilder RMA smoothing.",
    "SSL Channel: SMA(high,10) and SMA(low,10); state flips bullish when close is above high SMA and bearish when close is below low SMA.",
    "Primary LONG: SSL flips bullish, price crosses above EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 > 50.",
    "Primary SHORT: SSL flips bearish, price crosses below EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 < 50.",
    "20% reversal-skip filter: if the previous trade moved at least 20% in favor from entry, skip the opposite SSL flip entry on the exit candle.",
    "Initial stop: 1.5 ATR14 from entry.",
    "TP1: 2.5 ATR14 from entry; close 50% at TP1.",
    "Runner A: after TP1, move remaining 50% stop to breakeven and exit runner on opposite SSL flip or breakeven stop.",
    "Continuation entries are disabled.",
    "Risk-off overlay is disabled.",
    "Round-trip trading cost is included in R results.",
]


def build_docx(result: dict) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("NXT v3.2 Latest System And Indicators", 0)
    doc.add_paragraph("Current promoted latest system as of this workbook refresh.")

    doc.add_heading("Backtest Snapshot", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    for k, v in [
        ("System", result["systemVersion"]),
        ("Data standard", "Binance native 1D candles"),
        ("ATR standard", "TradingView ATR14 Wilder RMA"),
        ("Trades", result["stats"]["trades"]),
        ("Win rate", f"{result['stats']['winRate']:.2%}"),
        ("Total R", f"{result['stats']['totalR']:.2f}R"),
        ("Average R", f"{result['stats']['avgR']:.2f}R"),
        ("Max DD R", f"{result['stats']['maxDrawdownR']:.2f}R"),
        ("20K Account Ending", f"${20000 + result['stats']['totalR'] * 1000:,.2f}"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)

    doc.add_heading("Current Rules", level=1)
    for text in RULES:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("20% Filter Audit", level=1)
    doc.add_paragraph("Signals skipped by the promoted 20% reversal-skip filter:")
    skipped = result.get("skippedSignals", [])
    if skipped:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Symbol", "Date", "Side", "Reason"]):
            table.cell(0, i).text = h
        for item in skipped:
            row = table.add_row().cells
            row[0].text = item["symbol"].replace("USDT", "")
            row[1].text = item["date"]
            row[2].text = item["side"]
            row[3].text = item["reason"]
    else:
        doc.add_paragraph("No skipped signals in this run.")

    doc.save(LATEST_DOCX)


def fix_workbook_labels(path: Path) -> None:
    wb = load_workbook(path)
    if "Summary" in wb.sheetnames:
        ws = wb["Summary"]
        ws["A1"] = "NXT v3.2 Latest - TV ATR + 20% Reversal-Skip"
        ws["A2"] = "BTC/SOL/SUI 6Y | Runner A | 20% reversal-skip | no continuation | no risk-off | Binance native daily candles."
    if "Trades" in wb.sheetnames:
        wb["Trades"]["A1"] = "Detailed Trades - NXT v3.2 TV ATR + 20% Reversal-Skip"
    for sheet in ["BTC", "SOL", "SUI"]:
        if sheet in wb.sheetnames:
            wb[sheet]["A1"] = f"{sheet} - NXT v3.2 TV ATR + 20% Reversal-Skip"
    wb.save(path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    LATEST.mkdir(parents=True, exist_ok=True)

    all_trades = []
    all_skipped = []
    datasets = {}
    for symbol in base.SYMBOLS:
        candles = tv_atr.enrich_tv_atr(native.fetch_native_1d(symbol))
        datasets[symbol] = {
            "dailyRows": len(candles),
            "firstDay": candles[0]["localDate"],
            "lastDay": candles[-1]["localDate"],
            "source": "Binance spot native 1D klines",
        }
        trades, skipped = backtest_symbol(symbol, candles)
        all_trades.extend(trades)
        all_skipped.extend(skipped)

    all_trades.sort(key=lambda x: x["exitTime"])
    result = {
        "generatedAt": datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "systemVersion": "NXT v3.2 Simple + Binance Native 1D + TradingView ATR RMA + Runner A + 20% Reversal-Skip + No Continuation + No Risk-Off",
        "period": {
            "start": base.START_DATE.isoformat(),
            "end": (base.END_DATE - base.timedelta(days=1)).isoformat(),
            "timezone": "Binance native daily candles",
        },
        "symbols": base.SYMBOLS,
        "stats": base.stats(all_trades),
        "trades": all_trades,
        "skippedSignals": all_skipped,
        "datasets": datasets,
        "assumptions": RULES,
    }

    OUT_JSON.write_text(json.dumps(result, indent=2), encoding="utf-8")
    native.OUT_XLSX = OUT_XLSX
    native.build_workbook(result)
    fix_workbook_labels(OUT_XLSX)

    shutil.copy2(OUT_JSON, LATEST_JSON)
    shutil.copy2(OUT_XLSX, LATEST_XLSX)
    build_docx(result)
    LATEST_SUMMARY.write_text(
        "\n".join(
            [
                "# Latest NXT System",
                "",
                f"System: {result['systemVersion']}",
                "",
                f"Trades: {result['stats']['trades']}",
                f"Total R: {result['stats']['totalR']:.2f}R",
                f"Max DD R: {result['stats']['maxDrawdownR']:.2f}R",
                f"Win rate: {result['stats']['winRate']:.2%}",
                f"20K Account ending: ${20000 + result['stats']['totalR'] * 1000:,.2f}",
                f"Skipped by 20% filter: {len(all_skipped)}",
                "",
                "Notes: Uses Binance native 1D candles, TradingView ATR RMA, Runner A, and the 20% reversal-skip filter. Continuation and risk-off are disabled.",
                "",
                f"Workbook: {LATEST_XLSX.name}",
                f"JSON: {LATEST_JSON.name}",
                f"System doc: {LATEST_DOCX.name}",
            ]
        ),
        encoding="utf-8",
    )
    print(json.dumps({"latestJson": str(LATEST_JSON), "latestXlsx": str(LATEST_XLSX), "latestDocx": str(LATEST_DOCX), "stats": result["stats"], "skippedSignals": all_skipped}, indent=2))


if __name__ == "__main__":
    main()
