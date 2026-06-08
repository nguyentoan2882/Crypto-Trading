from __future__ import annotations

import json
import shutil
from copy import copy
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

import backtest_nxt31_utc7_latest as base


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "templates" / "NXT_Backtest_Workbook_Template.xlsx"
CACHE = ROOT / "data_cache" / "binance_spot_1d"
OUT_DIR = ROOT / "outputs" / "nxt33_native_1d_anti_reversal_6y"
OUT_JSON = OUT_DIR / "nxt33_native_1d_anti_reversal_6y_results.json"
OUT_XLSX = OUT_DIR / "NXT33_Native1D_AntiReversal_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST = ROOT / "latest"
ARCHIVE = ROOT / "outputs" / "archive_from_latest"
LATEST_JSON = LATEST / "NXT_Latest_NXT33_Native1D_AntiReversal_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.json"
LATEST_XLSX = LATEST / "NXT_Latest_NXT33_Native1D_AntiReversal_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST_DOCX = LATEST / "NXT_Latest_NXT33_System_And_Indicators.docx"
LATEST_SUMMARY = LATEST / "NXT_Latest_Summary.md"

SYMBOLS = base.SYMBOLS
START_DATE = base.START_DATE
END_DATE = base.END_DATE
WARMUP_DATE = base.WARMUP_DATE


def date_label(ms: int) -> str:
    return datetime.fromtimestamp(ms / 1000, timezone.utc).date().isoformat()


def fetch_native_1d(symbol: str) -> list[dict]:
    path = CACHE / f"{symbol}.json"
    rows = json.loads(path.read_text(encoding="utf-8"))
    out = []
    for r in rows:
        d = date_label(r["time"])
        if WARMUP_DATE <= date.fromisoformat(d) <= END_DATE:
            item = dict(r)
            item["localDate"] = d
            out.append(item)
    return out


def backtest_symbol(symbol: str, candles: list[dict]) -> list[dict]:
    trades, pos, n = [], None, 1
    last_profitable_runner_exit = None
    for i in range(55, len(candles) - 1):
        c, prev, nxt = candles[i], candles[i - 1], candles[i + 1]
        next_date = date.fromisoformat(nxt["localDate"])
        if next_date < START_DATE or next_date >= END_DATE:
            continue
        if pos:
            side = pos["side"]
            ssl_flip = (side == "LONG" and prev["ssl"] == 1 and c["ssl"] == -1) or (side == "SHORT" and prev["ssl"] == -1 and c["ssl"] == 1)
            exit_price = reason = None
            if side == "LONG":
                if c["low"] <= pos["stop"]:
                    exit_price = pos["stop"]
                    reason = "Breakeven stop" if pos["triggered"] else "Stop loss"
                else:
                    if not pos["triggered"] and c["high"] >= pos["tp"]:
                        pos["triggered"] = True
                        pos["tp1Time"] = c["localDate"]
                        pos["stop"] = pos["entry"]
                        pos["realizedR"] += 0.5 * ((pos["tp"] - pos["entry"]) / pos["risk"])
                    if ssl_flip:
                        exit_price = c["close"]
                        reason = "Runner exit: SSL bearish flip"
            else:
                if c["high"] >= pos["stop"]:
                    exit_price = pos["stop"]
                    reason = "Breakeven stop" if pos["triggered"] else "Stop loss"
                else:
                    if not pos["triggered"] and c["low"] <= pos["tp"]:
                        pos["triggered"] = True
                        pos["tp1Time"] = c["localDate"]
                        pos["stop"] = pos["entry"]
                        pos["realizedR"] += 0.5 * ((pos["entry"] - pos["tp"]) / pos["risk"])
                    if ssl_flip:
                        exit_price = c["close"]
                        reason = "Runner exit: SSL bullish flip"
            if exit_price is not None:
                rem = 0.5 if pos["triggered"] else 1.0
                rem_r = (exit_price - pos["entry"]) / pos["risk"] if side == "LONG" else (pos["entry"] - exit_price) / pos["risk"]
                gross = pos["realizedR"] + rem * rem_r
                net = gross - base.cost_r(pos["entry"], pos["risk"])
                trades.append({
                    "symbol": symbol,
                    "tradeNo": n,
                    "side": side,
                    "signalTime": pos["signalDate"],
                    "entryTime": pos["entryDate"],
                    "entryPrice": pos["entry"],
                    "initialStop": pos["initialStop"],
                    "finalStop": pos["stop"],
                    "riskPerUnit": pos["risk"],
                    "tp1": pos["tp"],
                    "tp1Time": pos["tp1Time"],
                    "exitTime": c["localDate"],
                    "exitPrice": exit_price,
                    "exitReason": reason,
                    "grossRMultiple": gross,
                    "costR": base.cost_r(pos["entry"], pos["risk"]),
                    "rMultiple": net,
                    "atr14": pos["atr14"],
                    "rsi14": pos["rsi14"],
                    "distanceToEma50Atr": pos["distance"],
                    "notes": "Binance native 1D candles; NXT v3.3 Runner A; anti-immediate-reversal; no continuation; no risk-off",
                })
                if net > 0 and reason.startswith("Runner exit"):
                    last_profitable_runner_exit = {"index": i, "side": side}
                n += 1
                pos = None
            if pos:
                continue

        if any(c[k] is None for k in ["ema20", "ema50", "atr14", "rsi14", "ssl"]) or prev["ssl"] is None:
            continue
        dist = abs(c["close"] - c["ema50"]) / c["atr14"]
        long_ok = prev["ssl"] == -1 and c["ssl"] == 1 and base.recent_cross(candles, i, "LONG") and dist <= 2 and c["rsi14"] > 50
        short_ok = prev["ssl"] == 1 and c["ssl"] == -1 and base.recent_cross(candles, i, "SHORT") and dist <= 2 and c["rsi14"] < 50
        if last_profitable_runner_exit and i - last_profitable_runner_exit["index"] <= 1:
            if long_ok and last_profitable_runner_exit["side"] == "SHORT":
                long_ok = False
            if short_ok and last_profitable_runner_exit["side"] == "LONG":
                short_ok = False
        if not (long_ok or short_ok):
            continue
        side = "LONG" if long_ok else "SHORT"
        risk = c["atr14"] * 1.5
        entry = nxt["open"]
        pos = {
            "side": side,
            "signalDate": c["localDate"],
            "entryDate": nxt["localDate"],
            "entry": entry,
            "initialStop": entry - risk if side == "LONG" else entry + risk,
            "stop": entry - risk if side == "LONG" else entry + risk,
            "risk": risk,
            "tp": entry + c["atr14"] * 2.5 if side == "LONG" else entry - c["atr14"] * 2.5,
            "triggered": False,
            "tp1Time": "",
            "realizedR": 0.0,
            "atr14": c["atr14"],
            "rsi14": c["rsi14"],
            "distance": dist,
        }
    return trades


def clear(ws, rows=700, cols=32):
    for row in ws.iter_rows(min_row=1, max_row=max(rows, ws.max_row), min_col=1, max_col=max(cols, ws.max_column)):
        for cell in row:
            if cell.__class__.__name__ != "MergedCell":
                cell.value = None


def rowset(ws, r, values):
    for c, v in enumerate(values, 1):
        ws.cell(r, c).value = v


def copy_layout(src, dst):
    dst.sheet_format.defaultColWidth = src.sheet_format.defaultColWidth
    dst.sheet_format.defaultRowHeight = src.sheet_format.defaultRowHeight
    dst.sheet_view.showGridLines = src.sheet_view.showGridLines
    for c in range(1, max(src.max_column, dst.max_column) + 1):
        letter = get_column_letter(c)
        if src.column_dimensions.get(letter) and src.column_dimensions[letter].width:
            dst.column_dimensions[letter].width = src.column_dimensions[letter].width
    for r in range(1, max(src.max_row, dst.max_row) + 1):
        sr = r if r <= src.max_row else 5
        if src.row_dimensions.get(sr) and src.row_dimensions[sr].height:
            dst.row_dimensions[r].height = src.row_dimensions[sr].height
        for c in range(1, max(src.max_column, dst.max_column) + 1):
            a = src.cell(sr, min(c, src.max_column))
            b = dst.cell(r, c)
            if a.has_style:
                b._style = copy(a._style)
            b.number_format = a.number_format
            b.alignment = copy(a.alignment)
            b.font = copy(a.font)
            b.fill = copy(a.fill)
            b.border = copy(a.border)


def build_workbook(result: dict):
    tpl = load_workbook(TEMPLATE)
    wb = load_workbook(TEMPLATE)
    for ws in wb.worksheets:
        clear(ws)
    ws = wb["Summary"]
    ws["A1"] = "NXT v3.3 Latest - Binance Native 1D"
    ws["A2"] = "BTC/SOL/SUI 6Y | Runner A | anti-immediate-reversal | no continuation | no risk-off | Binance native daily candles."
    rows = [
        ["Metric", "Value"],
        ["Trades", result["stats"]["trades"]],
        ["Win Rate", result["stats"]["winRate"]],
        ["Total R", result["stats"]["totalR"]],
        ["Average R", result["stats"]["avgR"]],
        ["Max DD R", result["stats"]["maxDrawdownR"]],
        ["Best R", result["stats"]["bestR"]],
        ["Worst R", result["stats"]["worstR"]],
        ["20K Account Ending", 20000 + result["stats"]["totalR"] * 1000],
    ]
    for i, row in enumerate(rows, 4):
        rowset(ws, i, row)

    headers = ["Symbol", "No", "Side", "Signal Date", "Entry Date", "Entry Price", "Initial Stop", "Final Stop", "Risk / Unit", "TP1", "TP1 Date", "Exit Date", "Exit Price", "Exit Reason", "R", "ATR14", "RSI14", "Distance EMA50 ATR", "Notes"]
    for sheet, subset in [
        ("Trades", result["trades"]),
        ("BTC", [t for t in result["trades"] if t["symbol"] == "BTCUSDT"]),
        ("SOL", [t for t in result["trades"] if t["symbol"] == "SOLUSDT"]),
        ("SUI", [t for t in result["trades"] if t["symbol"] == "SUIUSDT"]),
    ]:
        ws = wb[sheet]
        ws["A1"] = f"{sheet} - NXT v3.3 Binance Native 1D" if sheet != "Trades" else "Detailed Trades - NXT v3.3 Binance Native 1D"
        ws["A2"] = "One completed trade per row."
        rowset(ws, 4, headers)
        for i, t in enumerate(subset, 5):
            rowset(ws, i, [t["symbol"].replace("USDT", ""), t["tradeNo"], t["side"], t["signalTime"], t["entryTime"], t["entryPrice"], t["initialStop"], t["finalStop"], t["riskPerUnit"], t["tp1"], t["tp1Time"], t["exitTime"], t["exitPrice"], t["exitReason"], t["rMultiple"], t["atr14"], t["rsi14"], t["distanceToEma50Atr"], t["notes"]])

    ws = wb["20K Account"]
    ws["A1"] = "20K Account"
    ws["A2"] = "Assumes 1R = $1,000 on a $20,000 starting account."
    rowset(ws, 4, ["Starting Equity", "Total R", "Ending Equity"])
    rowset(ws, 5, [20000, result["stats"]["totalR"], 20000 + result["stats"]["totalR"] * 1000])

    ass = wb["Assumptions"]
    ass["A1"] = "Assumptions"
    rowset(ass, 4, ["#", "Assumption"])
    for i, line in enumerate(result["assumptions"], 5):
        rowset(ass, i, [i - 4, line])

    quality = wb["Data Quality"]
    quality["A1"] = "Data Quality"
    rowset(quality, 4, ["Symbol", "Daily Rows", "First Day", "Last Day", "Source"])
    for i, (sym, q) in enumerate(result["datasets"].items(), 5):
        rowset(quality, i, [sym.replace("USDT", ""), q["dailyRows"], q["firstDay"], q["lastDay"], q["source"]])

    for sheet in ["Summary", "Trades", "BTC", "SOL", "SUI", "Assumptions", "Data Quality", "20K Account", "Equity Curve"]:
        if sheet in wb.sheetnames and sheet in tpl.sheetnames:
            copy_layout(tpl[sheet], wb[sheet])
    wb.save(OUT_XLSX)


def build_doc(data: dict):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("NXT v3.3 Simple - Binance Native 1D", 0)
    doc.add_paragraph("Latest selected system: Binance native daily candles, Runner A, anti-immediate-reversal, no continuation module, no risk-off overlay.")
    doc.add_heading("Backtest Snapshot", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    for k, v in [
        ("System", data["systemVersion"]),
        ("Data standard", "Binance native 1D candles"),
        ("Trades", data["stats"]["trades"]),
        ("Win rate", f"{data['stats']['winRate']:.2%}"),
        ("Total R", f"{data['stats']['totalR']:.2f}R"),
        ("Average R", f"{data['stats']['avgR']:.2f}R"),
        ("Max DD R", f"{data['stats']['maxDrawdownR']:.2f}R"),
        ("20K Account Ending", f"${20000 + data['stats']['totalR'] * 1000:,.2f}"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
    doc.add_heading("Rules", level=1)
    for text in data["assumptions"]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(LATEST_DOCX)


def refresh_latest(result: dict):
    LATEST.mkdir(parents=True, exist_ok=True)
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    for item in LATEST.iterdir():
        if item.name.startswith("~$"):
            continue
        target = ARCHIVE / f"{timestamp}_{item.name}"
        try:
            shutil.move(str(item), str(target))
        except PermissionError as exc:
            if target.exists() and target.stat().st_size == 0:
                target.unlink()
            raise PermissionError(f"Cannot refresh latest because this file is open: {item}") from exc
    shutil.copy2(OUT_XLSX, LATEST_XLSX)
    shutil.copy2(OUT_JSON, LATEST_JSON)
    build_doc(result)
    LATEST_SUMMARY.write_text(
        "\n".join([
            "# Latest NXT System",
            "",
            f"System: {result['systemVersion']}",
            "",
            f"Trades: {result['stats']['trades']}",
            f"Total R: {result['stats']['totalR']:.2f}R",
            f"Max DD R: {result['stats']['maxDrawdownR']:.2f}R",
            f"Win rate: {result['stats']['winRate']:.2%}",
            f"20K Account ending: ${20000 + result['stats']['totalR'] * 1000:,.2f}",
            "",
            "Notes: Uses Binance native 1D candles to match TradingView Binance 1D indicator values. Anti-immediate-reversal is enabled. Continuation and risk-off are disabled.",
            "",
            f"Workbook: {LATEST_XLSX.name}",
            f"JSON: {LATEST_JSON.name}",
            f"System doc: {LATEST_DOCX.name}",
        ]),
        encoding="utf-8",
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    all_trades = []
    datasets = {}
    for symbol in SYMBOLS:
        candles = base.enrich(fetch_native_1d(symbol))
        datasets[symbol] = {"dailyRows": len(candles), "firstDay": candles[0]["localDate"], "lastDay": candles[-1]["localDate"], "source": "Binance spot native 1D klines"}
        all_trades.extend(backtest_symbol(symbol, candles))
    all_trades.sort(key=lambda x: x["exitTime"])
    result = {
        "generatedAt": datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "systemVersion": "NXT v3.3 Simple + Binance Native 1D + Runner A + Anti-Immediate-Reversal + No Continuation + No Risk-Off",
        "period": {"start": START_DATE.isoformat(), "end": (END_DATE - timedelta(days=1)).isoformat(), "timezone": "Binance native daily candles"},
        "symbols": SYMBOLS,
        "stats": base.stats(all_trades),
        "trades": all_trades,
        "datasets": datasets,
        "assumptions": [
            "Daily candles use Binance native 1D klines, matching TradingView BTCUSDT 1D Binance indicator values.",
            "Primary LONG: SSL flips bullish, EMA20 cross within last 3 candles, distance to EMA50 <= 2 ATR, RSI14 > 50.",
            "Primary SHORT: SSL flips bearish, EMA20 cross down within last 3 candles, distance to EMA50 <= 2 ATR, RSI14 < 50.",
            "Runner A: close 50% at TP1 = 2.5 ATR, move remaining 50% stop to breakeven, exit runner on opposite SSL flip or stop.",
            "Anti-immediate-reversal: after a profitable runner exit by opposite SSL flip, the system blocks an opposite-direction entry for the next 1 candle.",
            "Continuation module is disabled.",
            "Risk-off overlay is disabled.",
            "Cost model remains 0.06% fee and 0.05% slippage per side; funding is not included.",
        ],
    }
    OUT_JSON.write_text(json.dumps(result, indent=2), encoding="utf-8")
    build_workbook(result)
    refresh_latest(result)
    btc_march = [t for t in all_trades if t["symbol"] == "BTCUSDT" and ("2021-03" in t["signalTime"] or "2021-03" in t["entryTime"])]
    print(json.dumps({"outJson": str(OUT_JSON), "outXlsx": str(OUT_XLSX), "latestXlsx": str(LATEST_XLSX), "stats": result["stats"], "btcMarch2021": btc_march}, indent=2))


if __name__ == "__main__":
    main()
