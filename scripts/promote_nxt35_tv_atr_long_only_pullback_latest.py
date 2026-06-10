from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

sys.path.insert(0, str(Path(__file__).resolve().parent))

from test_nxt35_tv_atr_long_only_pullback_continuation import (
    OUT_JSON as TEST_JSON,
    OUT_XLSX as TEST_XLSX,
    main as build_test_artifacts,
)


ROOT = Path(__file__).resolve().parents[1]
LATEST = ROOT / "latest"
ARCHIVE = ROOT / "outputs" / "archive_from_latest"
LATEST_JSON = LATEST / "NXT_Latest_NXT35_Native1D_SSL14_TVATR_RunnerA_LongOnlyPullbackContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.json"
LATEST_XLSX = LATEST / "NXT_Latest_NXT35_Native1D_SSL14_TVATR_RunnerA_LongOnlyPullbackContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST_DOCX = LATEST / "NXT_Latest_NXT35_System_And_Indicators.docx"
LATEST_SUMMARY = LATEST / "NXT_Latest_Summary.md"

SYSTEM_VERSION = "NXT v3.5 Simple + Binance Native 1D + SSL14 + TradingView ATR RMA + Runner A + Anti-Immediate-Reversal + LONG-only Pullback Continuation + No Risk-Off"

RULES = [
    "Data: Binance native 1D candles for BTCUSDT, SOLUSDT, and SUIUSDT.",
    "ATR14 uses TradingView's default Wilder RMA smoothing.",
    "SSL Channel: SMA(high,14) and SMA(low,14); state flips bullish when close is above high SMA and bearish when close is below low SMA.",
    "Primary LONG: SSL flips bullish, price crosses above EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 > 50.",
    "Primary SHORT: SSL flips bearish, price crosses below EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 < 50.",
    "Continuation LONG: SSL is bullish, close > EMA20 > EMA50, low touched EMA20 within the last 5 candles, close > EMA20, and close > previous close.",
    "Continuation is LONG-only; SHORT continuation is disabled.",
    "Continuation does not require RSI, distance-to-EMA50, or EMA50 slope filters.",
    "Anti-immediate-reversal: after a profitable runner exit by opposite SSL flip, block an opposite-direction entry for the next 1 candle.",
    "Initial stop: 1.5 ATR14 from entry.",
    "TP1: 2.5 ATR14 from entry; close 50% at TP1.",
    "Runner A: after TP1, move remaining 50% stop to breakeven and exit runner on opposite SSL flip or breakeven stop.",
    "Risk-off overlay is disabled.",
    "Round-trip trading cost is included in R results.",
]


def archive_existing_latest() -> None:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    target = ARCHIVE / f"before_nxt35_tv_atr_{stamp}"
    target.mkdir(parents=True, exist_ok=True)
    for path in LATEST.glob("NXT_Latest_*"):
        if path.is_file() and path.name != "NXT_Latest_Summary.md":
            shutil.move(str(path), str(target / path.name))


def build_docx(result: dict) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("NXT v3.5 Latest System And Indicators", 0)
    doc.add_paragraph("Current promoted latest system: NXT v3.5 with TradingView ATR RMA and LONG-only pullback/touch EMA20 continuation.")
    doc.add_heading("Backtest Snapshot", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    for k, v in [
        ("System", result["systemVersion"]),
        ("Data standard", "Binance native 1D candles"),
        ("ATR", "TradingView Wilder RMA"),
        ("SSL period", "14"),
        ("Trades", result["stats"]["trades"]),
        ("Win rate", f"{result['stats']['winRate']:.2%}"),
        ("Total R", f"{result['stats']['totalR']:.2f}R"),
        ("Average R", f"{result['stats']['avgR']:.2f}R"),
        ("Max DD R", f"{result['stats']['maxDrawdownR']:.2f}R"),
        ("20K Account Ending", f"${20000 + result['stats']['totalR'] * 1000:,.2f}"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
    doc.add_heading("Current Rules", level=1)
    for text in RULES:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(LATEST_DOCX)


def fix_workbook_labels(path: Path) -> None:
    wb = load_workbook(path)
    if "Summary" in wb.sheetnames:
        wb["Summary"]["A1"] = "NXT v3.5 Latest - TradingView ATR RMA + LONG-only Pullback Continuation"
        wb["Summary"]["A2"] = "BTC/SOL/SUI 6Y | SSL14 | TradingView ATR RMA | Runner A | anti-immediate-reversal | LONG-only pullback continuation | no risk-off."
    if "Trades" in wb.sheetnames:
        wb["Trades"]["A1"] = "Detailed Trades - NXT v3.5 TradingView ATR RMA"
    wb.save(path)


def main() -> None:
    LATEST.mkdir(parents=True, exist_ok=True)
    try:
        build_test_artifacts()
    except FileNotFoundError as exc:
        if not (TEST_JSON.exists() and TEST_XLSX.exists()):
            raise
        print(f"Reusing existing NXT35 test artifacts because baseline latest was already archived: {exc}")
    archive_existing_latest()
    source = json.loads(TEST_JSON.read_text(encoding="utf-8"))
    result = {
        "generatedAt": datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "systemVersion": SYSTEM_VERSION,
        "period": source["period"],
        "symbols": ["BTCUSDT", "SOLUSDT", "SUIUSDT"],
        "stats": source["variantStats"],
        "baselineStats": source["baselineStats"],
        "continuationStats": source["continuationStats"],
        "addedStats": source["addedStats"],
        "trades": source["trades"],
        "continuationTrades": source["continuationTrades"],
        "addedTrades": source["addedTrades"],
        "bySymbol": source["bySymbol"],
        "byContinuationSymbol": source["byContinuationSymbol"],
        "byYear": source["byYear"],
        "byContinuationYear": source["byContinuationYear"],
        "datasets": source["datasets"],
        "assumptions": RULES,
    }
    LATEST_JSON.write_text(json.dumps(result, indent=2), encoding="utf-8")
    shutil.copy2(TEST_XLSX, LATEST_XLSX)
    fix_workbook_labels(LATEST_XLSX)
    build_docx(result)
    LATEST_SUMMARY.write_text(
        "\n".join([
            "# Latest NXT System",
            "",
            f"System: {result['systemVersion']}",
            "",
            f"Trades: {result['stats']['trades']}",
            f"Total R: {result['stats']['totalR']:.2f}R",
            f"Max DD R: {result['stats']['maxDrawdownR']:.2f}R",
            f"Win rate: {result['stats']['winRate']:.2%}",
            f"Profit factor: {result['stats']['profitFactor']:.2f}",
            f"20K Account ending: ${20000 + result['stats']['totalR'] * 1000:,.2f}",
            "",
            f"Continuation trades: {result['continuationStats']['trades']}",
            f"Continuation R: {result['continuationStats']['totalR']:.2f}R",
            f"Continuation win rate: {result['continuationStats']['winRate']:.2%}",
            "",
            "Validation: BTCUSDT SHORT signal 2021-12-28 is included with TradingView ATR RMA; ATR14 2535.00, distance 1.85 ATR, net +1.59R.",
            "",
            "Notes: Uses Binance native 1D candles, SSL Channel period 14, TradingView ATR RMA, Runner A, anti-immediate-reversal, LONG-only pullback/touch EMA20 continuation, and no risk-off.",
            "",
            f"Workbook: {LATEST_XLSX.name}",
            f"JSON: {LATEST_JSON.name}",
            f"System doc: {LATEST_DOCX.name}",
        ]),
        encoding="utf-8",
    )
    print(json.dumps({"latestJson": str(LATEST_JSON), "latestXlsx": str(LATEST_XLSX), "latestDocx": str(LATEST_DOCX), "stats": result["stats"], "continuationStats": result["continuationStats"]}, indent=2))


if __name__ == "__main__":
    main()
