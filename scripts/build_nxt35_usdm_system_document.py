from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LATEST = ROOT / "latest"
SOURCE_JSON = LATEST / "NXT_Latest_NXT35_USDM_BlockShortAfterLosingLong_FundingAdjusted_20K.json"
OUT = ROOT / "outputs" / "nxt35_block_short_after_losing_long_usdm" / "NXT35_USDM_BlockShortAfterLosingLong_Promoted_System_And_Indicators.docx"

NAVY = "1F4E79"
BLUE = "2F75B5"
PALE = "D9EAF7"
GRAY = "F2F2F2"


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, value: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, NAVY)
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        if widths:
            cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx == 0:
                shade(cells[idx], PALE)
            set_cell_text(cells[idx], value, bold=idx == 0)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)


def note(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "EAF2F8")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    a = p.add_run(f"{label}: ")
    a.bold = True
    b = p.add_run(text)
    for run in (a, b):
        run.font.name = "Aptos"
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_rule(doc: Document, rule_id: str, name: str, when: str, action: str, scope: str, detail: str) -> None:
    if rule_id in {"I-02", "P-04"}:
        doc.add_page_break()
    heading(doc, f"{rule_id} — {name}", 2)
    add_table(doc, ["Thuộc tính", "Quy định"], [
        ["Khi áp dụng", when],
        ["Hành động", action],
        ["Phạm vi", scope],
        ["Chi tiết kiểm soát", detail],
    ], [1.45, 5.75])


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Title", 22, NAVY), ("Heading 1", 15, NAVY), ("Heading 2", 11.5, BLUE), ("Heading 3", 10.5, NAVY)]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("NXT v3.5 USD-M — System Rulebook | Promoted latest").font.size = Pt(8)


def main() -> None:
    data = json.loads(SOURCE_JSON.read_text(encoding="utf-8"))
    stats = data["fundingAdjustedStats"]
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("NXT v3.5 USD-M System Rulebook")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("BTCUSDT · BNBUSDT · SOLUSDT | 1D | Promoted latest").bold = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph("Phiên bản vận hành: Block SHORT sau khi LONG trước TP1 thoát bằng SSL bearish flip với net R < 0.", style="Subtitle")
    add_table(doc, ["Mục", "Giá trị"], [
        ["Data / session", "Binance USD-M perpetual 1D; nến 00:00 UTC; chỉ dùng nến đã đóng."],
        ["Backtest", "2020-05-17 đến 2026-05-16; USD-M historical funding."],
        ["Kết quả promoted", f"{stats['trades']} trades | {stats['totalR']:.2f}R sau funding | PF {stats['profitFactor']:.2f} | max DD {stats['maxDrawdownR']:.2f}R."],
        ["Nguồn authority", "Latest JSON, latest regression và shared scanner core trong repository."],
    ], [1.55, 5.65])
    note(doc, "Nguyên tắc đọc", "Các rule bên dưới được sắp theo đúng thứ tự engine xử lý: đủ điều kiện dữ liệu → tín hiệu → quản lý vị thế → exit → block tín hiệu kế tiếp → tính kết quả.")

    heading(doc, "1. Phạm vi và thứ tự xử lý", 1)
    add_table(doc, ["Bước", "Engine làm gì"], [
        ["1", "Đọc nến 1D đã đóng, tính EMA20, EMA50, ATR14, RSI14 và SSL14."],
        ["2", "Nếu đang có vị thế: kiểm tra stop trước; sau đó TP1 / Early-BE; sau đó SSL exit."],
        ["3", "Khi vị thế đã đóng: ghi nhận các guard chống đảo chiều hoặc block SHORT mới."],
        ["4", "Đánh giá Primary LONG, Primary SHORT và LONG Continuation trên nến signal."],
        ["5", "Áp dụng các block. Nếu còn tín hiệu hợp lệ, entry tại open của nến ngày kế tiếp."],
        ["6", "Sau khi trade đóng: tính gross R, transaction cost R, rồi funding R để ra net R after funding."],
    ], [0.6, 6.6])

    heading(doc, "2. Data, thời điểm và chỉ báo", 1)
    add_rule(doc, "D-01", "Data và candle completeness", "Mỗi lượt scan/backtest.", "Chỉ xét candle USD-M 1D đã closed.", "BTCUSDT, BNBUSDT, SOLUSDT; UTC 00:00 session.", "Signal được xác nhận bằng close của candle signal. Entry dùng open của candle kế tiếp; không dùng nến đang chạy để tạo signal.")
    add_rule(doc, "I-01", "EMA20 và EMA50", "Tính trên close của 1D candle.", "Dùng trong trend, pullback và distance filter.", "Tất cả signal.", "EMA khởi tạo bằng SMA của period tương ứng, rồi cập nhật với alpha = 2/(period+1).")
    add_rule(doc, "I-02", "ATR14", "Tính từ True Range trong 14 candles.", "Dùng chuẩn hóa distance, stop và TP1.", "Tất cả signal/vị thế.", "TR = max(High-Low, |High-prev Close|, |Low-prev Close|). ATR14 của rulebook này là SMA(TR, 14), không phải ATR Wilder/RMA.")
    add_rule(doc, "I-03", "RSI14", "Tính trên close.", "Lọc momentum Primary.", "Primary LONG/SHORT.", "RSI dùng average gain/loss 14 kỳ theo smoothing Wilder sau seed ban đầu.")
    add_rule(doc, "I-04", "SSL14", "SMA14 của High và Low; duy trì state.", "Xác định flip bullish/bearish và exit SSL.", "Tất cả signal/vị thế.", "SSL = bullish (+1) khi Close > SMA14(High); bearish (-1) khi Close < SMA14(Low); nếu nằm giữa hai band thì giữ state trước đó.")

    heading(doc, "3. Điều kiện entry", 1)
    add_rule(doc, "E-00", "Điều kiện chung", "Trước khi xét từng setup.", "Từ chối signal nếu indicator chưa đủ dữ liệu.", "Tất cả signal.", "EMA20, EMA50, ATR14, RSI14, SSL hiện tại và SSL candle trước phải có giá trị. Distance = |Close − EMA50| / ATR14 phải <= 2.00 cho Primary.")
    add_rule(doc, "E-01", "Primary LONG", "SSL flip từ -1 sang +1 trên candle signal.", "Cho phép LONG Primary.", "Chỉ LONG Primary.", "Đồng thời phải có EMA20 cross-up trong candle signal hoặc 2 candles trước đó; RSI14 > 50; và distance-to-EMA50 <= 2 ATR.")
    add_rule(doc, "E-02", "Primary SHORT", "SSL flip từ +1 sang -1 trên candle signal.", "Cho phép SHORT Primary.", "Chỉ SHORT Primary.", "Đồng thời phải có EMA20 cross-down trong candle signal hoặc 2 candles trước đó; RSI14 < 50; và distance-to-EMA50 <= 2 ATR.")
    add_rule(doc, "E-03", "LONG Continuation", "SSL flip từ -1 sang +1 và bối cảnh trend tăng.", "Cho phép LONG Continuation.", "Chỉ LONG Continuation; SHORT Continuation tắt.", "Close > EMA20 > EMA50. Trong 5 candles gần nhất phải có ít nhất một Low <= EMA20; candle signal phải Close > EMA20 và Close > close của candle trước.")
    add_rule(doc, "E-04", "Ưu tiên setup và entry", "Sau khi các entry rule đều đúng.", "Nếu LONG Primary và LONG Continuation cùng đúng, ghi nhận Primary.", "Tất cả entry.", "Side LONG được ưu tiên khi có LONG signal; signalType = Continuation chỉ khi Continuation đúng nhưng Primary LONG không đúng. Entry price = open của candle kế tiếp.")

    heading(doc, "4. Quản lý vị thế và exits", 1)
    add_rule(doc, "P-01", "Initial risk và stop", "Ngay khi entry.", "Đặt stop ban đầu và định nghĩa 1R.", "LONG/SHORT.", "Risk per unit = 1.5 × ATR14 của candle signal. LONG stop = Entry − Risk; SHORT stop = Entry + Risk.")
    add_rule(doc, "P-02", "TP1", "Khi giá chạm TP1 trước stop.", "Đóng 50% vị thế; runner còn 50%; stop chuyển về entry.", "LONG/SHORT.", "TP1 LONG = Entry + 2.5 × ATR14 signal. TP1 SHORT = Entry − 2.5 × ATR14 signal. Phần TP1 đóng tạo gross +0.8333R (50% × 2.5/1.5).")
    add_rule(doc, "P-03", "Early-BE 7%", "Từ candle đầu tiên sau entry, khi chưa TP1 và chưa Early-BE.", "Dời stop toàn bộ về entry cho các candles sau.", "LONG/SHORT.", "LONG: High >= Entry × 1.07. SHORT: Low <= Entry × 0.93. Stop được kiểm tra trước trigger trong cùng candle, nên Early-BE chỉ bảo vệ từ candle kế tiếp.")
    add_rule(doc, "P-04", "Stop priority", "Mỗi candle đang giữ vị thế.", "Kiểm tra stop trước TP1, Early-BE và SSL flip.", "LONG/SHORT.", "LONG bị stop nếu Low <= stop; SHORT bị stop nếu High >= stop. Khi stop là entry sau TP1/Early-BE, exit reason là Breakeven stop; nếu không là Stop loss.")
    add_rule(doc, "P-05", "SSL runner exit", "Sau khi stop không bị chạm trong candle.", "Đóng phần vị thế còn lại tại candle close khi SSL đảo chiều.", "LONG/SHORT.", "LONG exit khi SSL +1 → -1; SHORT exit khi SSL -1 → +1. Nếu TP1 và SSL exit cùng candle, TP1 được ghi trước rồi runner đóng ở close candle đó.")
    add_rule(doc, "P-06", "Tính R khi đóng trade", "Khi có stop hoặc SSL exit.", "Tính gross và net trước funding.", "LONG/SHORT.", "Remaining fraction = 1.0 khi chưa TP1, hoặc 0.5 sau TP1. Gross R = realized TP1 R + remaining fraction × price movement / Risk. Net R before funding = Gross R − transaction cost R.")

    heading(doc, "5. Signal blocks và anti-chop guards", 1)
    add_rule(doc, "G-01", "Anti-immediate-reversal sau profitable SSL exit", "Trade vừa đóng bởi SSL runner exit với net R before funding >= +0.50R.", "Block entry ngược chiều trên candle exit và candle kế tiếp.", "Primary + LONG Continuation theo hướng ngược.", "Nếu LONG vừa SSL exit có net >= +0.50R, block SHORT Primary. Nếu SHORT vừa SSL exit có net >= +0.50R, block LONG Primary và LONG Continuation. Window được tính i − exitIndex <= 1.")
    add_rule(doc, "G-02", "Block SHORT sau losing pre-TP1 LONG SSL exit", "Prior LONG chưa đạt TP1, thoát đúng bằng SSL bearish flip và net R before funding < 0.", "Block SHORT Primary trên candle exit và đúng 1 candle kế tiếp.", "Chỉ SHORT Primary; không block LONG setup.", "Rule được ghi sau khi tính net R của LONG exit. Điều kiện TP1 là false tại thời điểm exit. Đây là rule promoted mới nhằm tránh follow-through SHORT ngay sau khi LONG failure đã xác nhận bearish flip nhưng chất lượng tiếp diễn thấp.")
    note(doc, "Không có rule ATR Expansion Guard", "ATR14/SMA(ATR14,10) không được dùng làm entry filter trong bản latest này. Các threshold thử nghiệm trước đây không phải rule promoted.")

    heading(doc, "6. Cost, funding và báo cáo hiệu năng", 1)
    add_rule(doc, "C-01", "Transaction cost", "Mọi trade đóng.", "Trừ transaction cost khỏi gross R.", "Tất cả trades.", "Fee giả định 0.0006 mỗi chiều và slippage 0.0005 mỗi chiều; round-trip rate = 2 × (0.0006 + 0.0005) = 0.0022. Cost R = Entry × 0.0022 / Risk per unit.")
    add_rule(doc, "C-02", "USD-M funding", "Sau khi core trade được tính.", "Cộng/trừ funding R theo dữ liệu funding lịch sử.", "Tất cả trades trong backtest.", "Funding áp dụng từ entry đến hết exit date. Notional proxy = Entry/Risk. Fraction = 100% trước TP1 và 50% từ funding event tại/sau TP1. net R after funding = rMultiple + fundingR.")
    add_rule(doc, "C-03", "Portfolio reporting", "Kết thúc backtest.", "Tổng hợp theo thứ tự exit trade.", "Portfolio 3 symbols.", "Kết quả promoted: 235 trades, 163.90R after funding, win rate 47.66%, PF 2.97, max drawdown -7.23R. Rủi ro dollar trong workbook dùng 1R = $1,000 và starting equity $20,000.")

    heading(doc, "7. Pseudocode kiểm soát", 1)
    pseudo = [
        "for each closed 1D candle c:",
        "  if position: stop check → TP1/Early-BE → opposite SSL exit",
        "  on exit: record G-01 / G-02 state from net R before funding",
        "  evaluate E-01, E-02, E-03 on c",
        "  apply G-01 then G-02 blocks",
        "  if a signal remains: enter at next candle open with P-01 parameters",
        "  after exit: calculate C-01 and C-02 for reporting",
    ]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "1E293B")
    cell.text = ""
    for index, line in enumerate(pseudo):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph()
    note(doc, "Scope guard", "System document mô tả đúng rule family đã promote và không phải instruction để auto-execute lệnh. Live execution vẫn cần approval, exchange filters, account-risk controls và reconciliation riêng.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "NXT v3.5 USD-M System Rulebook"
    doc.core_properties.subject = "Promoted system rules and execution logic"
    doc.core_properties.author = "NXT"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
