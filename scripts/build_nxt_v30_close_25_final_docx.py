from __future__ import annotations

import json
import shutil
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Workspace\Codex\Crypto trading")
SOURCE_DIR = ROOT / "outputs" / "nxt_crypto_btc_sol_sui_6y_v23_close_25"
SOURCE_JSON = SOURCE_DIR / "nxt_v23_close_25_6y_results.json"
SOURCE_XLSX = SOURCE_DIR / "NXT_V23_Close_25_ATR_6Y_BTC_SOL_SUI.xlsx"

OUT_DIR = ROOT / "outputs" / "nxt_crypto_btc_sol_sui_6y_v30_close_25"
JSON_PATH = OUT_DIR / "nxt_v30_close_25_6y_results.json"
XLSX_PATH = OUT_DIR / "NXT_V30_Close_25_ATR_6Y_BTC_SOL_SUI.xlsx"
DOCX_PATH = OUT_DIR / "NXT_Trading_System_v3.0_Final.docx"


def read_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt_r(value: float) -> str:
    return f"{value:+.2f}R"


def fmt_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill: str = "EAF0F6") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 41, 55)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(17, 24, 39)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def summarize(data: dict) -> dict:
    trades = sorted(data["trades"], key=lambda t: t["exitTime"])
    total = sum(t["rMultiple"] for t in trades)
    wins = sum(1 for t in trades if t["rMultiple"] > 0)
    equity = peak = 0.0
    max_dd = 0.0
    for trade in trades:
        equity += trade["rMultiple"]
        peak = max(peak, equity)
        max_dd = min(max_dd, equity - peak)
    return {
        "trades": len(trades),
        "wins": wins,
        "losses": len(trades) - wins,
        "winRate": wins / len(trades),
        "totalR": total,
        "avgR": total / len(trades),
        "bestR": max(t["rMultiple"] for t in trades),
        "worstR": min(t["rMultiple"] for t in trades),
        "maxDD": max_dd,
        "tpRate": sum(1 for t in trades if t.get("tp1Hit") == "Yes") / len(trades),
        "exitReasons": Counter(t["exitReason"] for t in trades),
    }


def compound_projection(trades: list[dict], start_equity: float = 50_000, risk: float = 0.02) -> tuple[float, float]:
    equity = peak = start_equity
    max_dd = 0.0
    for trade in sorted(trades, key=lambda t: t["exitTime"]):
        equity *= 1 + risk * trade["rMultiple"]
        peak = max(peak, equity)
        max_dd = min(max_dd, equity / peak - 1)
    return equity, max_dd


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12.5)

    header = section.header.paragraphs[0]
    header.text = "NXT Trading System v3.0"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(107, 114, 128)
    return doc


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_JSON, JSON_PATH)
    shutil.copy2(SOURCE_XLSX, XLSX_PATH)

    data = read_json(JSON_PATH)
    data["systemVersion"] = "NXT v3.0"
    data["versionChange"] = "v3.0 keeps NXT v2.3 entry/filter logic and changes exit management to a full close at 2.5 ATR."
    JSON_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

    overall = summarize(data)
    final_equity, compound_dd = compound_projection(data["trades"])
    doc = setup_doc()

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("NXT Trading System v3.0").bold = True
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Bản cập nhật rule: chốt toàn bộ vị thế tại 2.5 ATR")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(75, 85, 99)

    created = datetime.now(timezone.utc).astimezone().strftime("%d-%b-%Y %H:%M")
    add_table(
        doc,
        ["Hạng mục", "Giá trị"],
        [
            ["Version", "NXT v3.0"],
            ["Thay đổi chính", "Giữ entry/filter của v2.3; exit đổi sang chốt toàn bộ tại 2.5 ATR."],
            ["Dữ liệu backtest", f'{data["period"]["start"][:10]} đến {data["period"]["end"][:10]}'],
            ["Danh mục", "BTCUSDT, SOLUSDT, SUIUSDT"],
            ["Nguồn dữ liệu", data["source"]],
            ["Workbook chi tiết", str(XLSX_PATH)],
            ["Ngày xuất tài liệu", created],
        ],
        [4.2, 11.8],
    )

    add_heading(doc, "Rule v3.0", 1)
    add_table(
        doc,
        ["Nhóm rule", "Thiết lập v3.0"],
        [
            ["Timeframe", "Daily signal, entry tại open của nến daily kế tiếp."],
            ["Weekly regime", "OFF, không dùng làm filter entry."],
            ["Trend trigger", "SSL Channel 10/10 crossover, xấp xỉ bằng SMA(high,10) và SMA(low,10)."],
            ["EMA confirmation", "Giá cross EMA20 trong vòng 3 nến gần nhất theo hướng lệnh."],
            ["EMA50 distance", "Khoảng cách từ close tới EMA50 <= 2 ATR(14)."],
            ["Volume filter", "Net Volume Binance = taker buy base volume x 2 - total volume. Long cần > 0, Short cần < 0."],
            ["Stop loss", "1.5 ATR(14) từ entry."],
            ["Take profit", "Chốt 100% vị thế tại 2.5 ATR(14)."],
            ["Exit khác", "Opposite SSL flip, stop loss, hoặc mark-to-market cuối kỳ nếu chưa chạm 2.5 ATR."],
            ["Cost model", "0.06% fee + 0.05% slippage mỗi chiều, trừ trực tiếp vào R từng lệnh."],
            ["Intraday assumption", "Nếu cùng ngày chạm stop và TP, ưu tiên stop trước TP để bảo thủ."],
        ],
        [4.4, 11.6],
    )

    add_heading(doc, "Kết quả backtest 6 năm", 1)
    add_table(
        doc,
        ["Metric", "Kết quả"],
        [
            ["Tổng số lệnh", str(overall["trades"])],
            ["Win / Loss", f'{overall["wins"]} / {overall["losses"]}'],
            ["Win rate", fmt_pct(overall["winRate"])],
            ["Total R", fmt_r(overall["totalR"])],
            ["Average R / trade", fmt_r(overall["avgR"])],
            ["Best / Worst trade", f'{fmt_r(overall["bestR"])} / {fmt_r(overall["worstR"])}'],
            ["Max drawdown", fmt_r(overall["maxDD"])],
            ["Tỷ lệ chạm TP 2.5 ATR", fmt_pct(overall["tpRate"])],
            ["Compound minh họa 50k, risk 2%", f"${final_equity:,.0f} | Max DD {compound_dd * 100:.2f}%"],
        ],
        [5.2, 5.8],
    )

    add_heading(doc, "Đóng góp theo coin", 2)
    rows = []
    for item in data["summary"]:
        rows.append([
            item["symbol"].replace("USDT", ""),
            str(item["trades"]),
            f'{item["wins"]}/{item["losses"]}',
            fmt_pct(item["winRate"]),
            fmt_r(item["totalR"]),
            fmt_r(item["avgR"]),
        ])
    add_table(doc, ["Coin", "Trades", "W/L", "Win rate", "Total R", "Avg R"], rows, [1.8, 1.7, 1.7, 2.2, 2.2, 2.2])

    add_heading(doc, "Lý do nâng version", 1)
    for item in [
        "TP1/TP2 gốc của v2.3 chốt 50% tại 1.5 ATR và 50% còn lại tại 2.5 ATR. Grid test cho thấy chốt sớm tại 1.5 ATR làm giảm expectancy.",
        "Khi quét TP1 từ 1.5 đến 2.5 ATR, mức 2.5 ATR cho Total R cao nhất trong mẫu 3 năm.",
        "Khi cố định TP1 2.5 ATR và quét TP2 lớn hơn, TP2 gần nhất 2.6 ATR là tốt nhất nhưng vẫn thấp hơn phương án chốt toàn bộ tại 2.5 ATR.",
        "Pareto 80/20 runner chỉ nhỉnh hơn bản gốc nhưng kém xa exit cố định 2.5 ATR, nên không chọn làm rule chính.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Data coverage", 1)
    coverage_rows = []
    for symbol, quality in data["datasets"].items():
        coverage_rows.append([
            symbol.replace("USDT", ""),
            str(quality["dailyCount"]),
            quality["firstDaily"][:10],
            quality["lastDaily"][:10],
            quality["source"],
        ])
    add_table(doc, ["Coin", "Daily candles", "First daily", "Last daily", "Source"], coverage_rows, [1.8, 2.4, 2.4, 2.4, 7.0])

    add_heading(doc, "Ghi chú triển khai", 1)
    for item in [
        "Backtest tính theo R và đã trừ fee/slippage, nhưng chưa tính funding, borrow cost, thuế, liquidation, minimum order, hoặc lỗi khớp lệnh khi live.",
        "Nếu risk 2% mỗi lệnh, cần thêm giới hạn tổng risk đang mở. Trong lịch sử 6 năm có 200 ngày có từ 2 coin trở lên đang mở cùng lúc, và 33 ngày có đủ 3 coin.",
        "Khuyến nghị thực chiến: mỗi lệnh tối đa 2% risk, tổng risk đang mở tối đa 4-6% tùy mức chịu drawdown.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Artifacts", 1)
    add_table(
        doc,
        ["Artifact", "Đường dẫn"],
        [
            ["Workbook v3.0", str(XLSX_PATH)],
            ["JSON v3.0", str(JSON_PATH)],
            ["Script 6Y close 2.5 ATR", str(ROOT / "scripts" / "backtest_nxt_v23_close_25_6y.mjs")],
        ],
        [4.0, 12.0],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
