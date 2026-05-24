from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(r"D:\Workspace\Codex\Crypto trading")
SRC_JSON = ROOT / "outputs" / "nxt31_utc7_latest_6y" / "nxt31_utc7_latest_6y_results.json"
SRC_XLSX = ROOT / "outputs" / "nxt31_utc7_latest_6y" / "NXT31_UTC7_Latest_RunnerA_RiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
LATEST = ROOT / "latest"
OUT_JSON = LATEST / "NXT_Latest_NXT32_UTC7_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.json"
OUT_XLSX = LATEST / "NXT_Latest_NXT32_UTC7_RunnerA_NoContinuation_NoRiskOff_6Y_BTC_SOL_SUI_20K.xlsx"
OUT_DOCX = LATEST / "NXT_Latest_NXT32_System_And_Indicators.docx"
SUMMARY = LATEST / "NXT_Latest_Summary.md"


def build_doc(data):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("NXT v3.2 Simple", 0)
    doc.add_paragraph("Latest selected system: UTC+7 daily candles, Runner A, no continuation module, no risk-off overlay.")

    doc.add_heading("Backtest Snapshot", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    for k, v in [
        ("System", "NXT v3.2 Simple"),
        ("Timezone", "Asia/Saigon UTC+7 daily candles"),
        ("Trades", data["stats"]["trades"]),
        ("Win rate", f"{data['stats']['winRate']:.2%}"),
        ("Total R", f"{data['stats']['totalR']:.2f}R"),
        ("Average R", f"{data['stats']['avgR']:.2f}R"),
        ("Max DD R", f"{data['stats']['maxDrawdownR']:.2f}R"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)

    doc.add_heading("Rules", level=1)
    for text in [
        "Daily candles are resampled to UTC+7 to match the local TradingView chart.",
        "Primary LONG: SSL flips bullish, EMA20 cross within last 3 candles, distance to EMA50 <= 2 ATR, RSI14 > 50.",
        "Primary SHORT: SSL flips bearish, EMA20 cross down within last 3 candles, distance to EMA50 <= 2 ATR, RSI14 < 50.",
        "Runner A: close 50% at 2.5 ATR, move remaining 50% to breakeven, exit runner on opposite SSL flip.",
        "Continuation module is disabled.",
        "Risk-off overlay is disabled.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(OUT_DOCX)


def main():
    LATEST.mkdir(parents=True, exist_ok=True)
    data = json.loads(SRC_JSON.read_text(encoding="utf-8"))
    simple = {
        "systemVersion": "NXT v3.2 Simple + UTC7 + Runner A + No Continuation + No Risk-Off",
        "source": str(SRC_JSON),
        "stats": data["stats"],
        "trades": data["trades"],
        "datasets": data["datasets"],
        "assumptions": [
            "UTC+7 daily candles.",
            "Runner A exit.",
            "Continuation disabled.",
            "Risk-off disabled.",
        ],
    }
    OUT_JSON.write_text(json.dumps(simple, indent=2), encoding="utf-8")
    shutil.copy2(SRC_XLSX, OUT_XLSX)
    build_doc(simple)
    SUMMARY.write_text(
        "\n".join(
            [
                "# Latest NXT System",
                "",
                "System: NXT v3.2 Simple + UTC+7 + Runner A + No Continuation + No Risk-Off",
                "",
                f"Total R: {data['stats']['totalR']:.2f}R",
                f"Max DD R: {data['stats']['maxDrawdownR']:.2f}R",
                f"Trades: {data['stats']['trades']}",
                f"Win rate: {data['stats']['winRate']:.2%}",
                "",
                "Notes: Continuation module and risk-off are disabled. UTC+7 daily candles are used to match the local TradingView setup.",
                "",
                f"Workbook: {OUT_XLSX.name}",
                f"JSON: {OUT_JSON.name}",
                f"System doc: {OUT_DOCX.name}",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
