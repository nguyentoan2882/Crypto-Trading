from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Workspace\Codex\Crypto trading")
LATEST = ROOT / "latest"
SOURCE = LATEST / "NXT_Latest_NXT31_RunnerA_RiskOff_6Y_BTC_SOL_SUI_20K.json"
OUT = LATEST / "NXT_Latest_NXT31_System_And_Indicators.docx"

ACCENT = "17324D"
LIGHT = "EAF0F6"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header=True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
            if header and r_idx == 0:
                set_cell_shading(cell, ACCENT)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(ACCENT)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def fmt_pct(x):
    return f"{x:.2%}"


def build():
    data = json.loads(SOURCE.read_text(encoding="utf-8"))
    stats = data["stats"]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(21)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    header = sec.header.paragraphs[0]
    header.text = "NXT v3.1 Latest System Specification"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    title = doc.add_paragraph(style="Title")
    title.add_run("NXT v3.1 + Runner A + Risk-Off")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Latest system and indicator specification for BTC/SOL/SUI daily crypto backtest").italic = True

    meta = doc.add_table(rows=5, cols=2)
    meta.cell(0, 0).text = "System version"
    meta.cell(0, 1).text = data["systemVersion"]
    meta.cell(1, 0).text = "Universe"
    meta.cell(1, 1).text = "BTCUSDT, SOLUSDT, SUIUSDT"
    meta.cell(2, 0).text = "Backtest period"
    meta.cell(2, 1).text = f"{data['period']['start'][:10]} to {data['period']['end'][:10]}"
    meta.cell(3, 0).text = "Account assumption"
    meta.cell(3, 1).text = "20,000 USD, 2.0% risk per trade, R-based accounting"
    meta.cell(4, 0).text = "Latest artifact folder"
    meta.cell(4, 1).text = str(LATEST)
    for row in meta.rows:
        set_cell_shading(row.cells[0], LIGHT)
        row.cells[0].paragraphs[0].runs[0].bold = True
    style_table(meta, header=False)

    add_heading(doc, "1. Latest Backtest Snapshot", 1)
    snap = doc.add_table(rows=1, cols=6)
    for i, h in enumerate(["Total R", "Max DD R", "Trades", "Win Rate", "Risk-Off Trades", "Final 20K Equity"]):
        snap.cell(0, i).text = h
    row = snap.add_row().cells
    values = [
        f"{stats['totalR']:.2f}R",
        f"{stats['maxDrawdownR']:.2f}R",
        stats["trades"],
        fmt_pct(stats["winRate"]),
        stats["riskOffTrades"],
        f"${stats['finalEquity20K']:,.0f}",
    ]
    for i, v in enumerate(values):
        row[i].text = str(v)
    style_table(snap)

    add_heading(doc, "2. TradingView Indicators", 1)
    ind = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Indicator", "Setting", "Usage", "TradingView note"]):
        ind.cell(0, i).text = h
    rows = [
        ["SSL Channel", "Length 10/10, SMA(high,10), SMA(low,10)", "Main direction flip signal", "Use project SSL Channel indicator."],
        ["EMA", "20 / 50 / 100 / 200 visible", "NXT v3.1 entry uses EMA20 cross and EMA50 distance", "Add EMA set or Moving Average indicators."],
        ["ATR", "ATR(14)", "Stop, TP trigger, distance filter", "Add ATR indicator, length 14."],
        ["RSI", "RSI(14), regime 50/50", "Noise filter: LONG > 50, SHORT < 50", "Use default RSI; middle band 50 is enough."],
        ["MACD", "Optional display only", "Not part of latest NXT v3.1 rule", "Can hide; previous MACD filter was rejected."],
    ]
    for r in rows:
        cells = ind.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v
    style_table(ind)

    add_heading(doc, "3. Entry Rules", 1)
    add_heading(doc, "LONG setup", 2)
    for text in [
        "SSL Channel flips bullish: previous SSL state bearish, current SSL state bullish.",
        "Price crosses above EMA20 within the last 3 daily candles.",
        "Distance from signal close to EMA50 is <= 2.0 ATR(14).",
        "RSI14 > 50 at signal candle close.",
        "Entry is next daily open after the signal candle closes.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "SHORT setup", 2)
    for text in [
        "SSL Channel flips bearish: previous SSL state bullish, current SSL state bearish.",
        "Price crosses below EMA20 within the last 3 daily candles.",
        "Distance from signal close to EMA50 is <= 2.0 ATR(14).",
        "RSI14 < 50 at signal candle close.",
        "Entry is next daily open after the signal candle closes.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4. Stop, Runner A, And Exit Rules", 1)
    for text in [
        "Initial stop is 1.5 x ATR(14) from entry price.",
        "Runner trigger is 2.5 x ATR(14) from entry.",
        "At 2.5 ATR trigger, close 50% of the position and move stop on the remaining 50% to breakeven.",
        "The remaining 50% runner exits on opposite SSL flip, breakeven stop, or end-of-test mark-to-market.",
        "Conservative candle handling is used: stop is checked before profit trigger when multiple levels may be touched in one daily candle.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "5. Risk-Off Overlay", 1)
    risk = data["riskOffRule"]
    for text in [
        "Risk-off changes position size only. It does not change entry signals or exit signals.",
        f"If the closed-trade equity curve is at or below {risk['triggerDrawdownR']:.0f}R from peak before a trade, count that trade at {risk['sizeMultiplierWhenTriggered']:.0%} size.",
        "If equity drawdown is above the trigger level, count the trade at 100% size.",
        "This overlay reduced Runner A max drawdown from -10.63R to -6.82R while keeping Total R at +63.90R.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "6. Implementation Notes", 1)
    notes = doc.add_table(rows=1, cols=2)
    notes.cell(0, 0).text = "Topic"
    notes.cell(0, 1).text = "Latest decision"
    decision_rows = [
        ["Net Volume", "Removed from latest system; RSI regime 50/50 replaced it."],
        ["MACD", "Tested as directional filter and rejected because it reduced total R too much."],
        ["RSI Variant", "RSI 50/50 selected for v3.1. RSI 52/48 reduces drawdown but sacrifices too much Total R when used as a hard filter."],
        ["Runner", "Runner A selected over Runner B because it captured more bull-run upside."],
        ["Risk-off", "Selected as current compromise because it keeps Total R within 10% reduction target better than hard filters."],
        ["Workbook template", "All future backtest outputs should use NXT_Backtest_Workbook_Template.xlsx format."],
    ]
    for r in decision_rows:
        cells = notes.add_row().cells
        cells[0].text = r[0]
        cells[1].text = r[1]
    style_table(notes)

    add_heading(doc, "7. TradingView Checklist", 1)
    for text in [
        "Chart: Binance BTCUSDT/SOLUSDT/SUIUSDT, timeframe 1D.",
        "Add SSL Channel and keep it visible.",
        "Add EMA 20/50/100/200 for visual context; EMA20 and EMA50 are required by the system.",
        "Add ATR(14), usually hidden after confirming values.",
        "Add RSI(14). Keep middle band 50 visible; LONG requires RSI above 50, SHORT below 50.",
        "MACD and Net Volume are not required for the latest system.",
    ]:
        add_number(doc, text)

    doc.add_page_break()
    add_heading(doc, "Appendix - Per Symbol Summary", 1)
    sym = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["Symbol", "Trades", "Wins", "Losses", "Win Rate", "Total R", "Avg R"]):
        sym.cell(0, i).text = h
    for row_data in data["summary"]:
        cells = sym.add_row().cells
        vals = [
            row_data["symbol"].replace("USDT", ""),
            row_data["trades"],
            row_data["wins"],
            row_data["losses"],
            fmt_pct(row_data["winRate"]),
            f"{row_data['totalR']:.2f}R",
            f"{row_data['avgR']:.2f}R",
        ]
        for i, v in enumerate(vals):
            cells[i].text = str(v)
    style_table(sym)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
