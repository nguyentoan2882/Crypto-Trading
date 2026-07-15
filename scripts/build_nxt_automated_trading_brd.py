from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "BRD_NXT_Automated_Signal_Trading_System_v0.2.docx"

BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "172B4D"
GREEN = "E8F3EC"
AMBER = "FFF4D6"
RED = "FDECEC"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    """Keep a table row intact when it reaches a page boundary."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=DARK, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run(run, size=9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 24, DARK, 0, 6),
        ("Subtitle", 13, MID_GRAY, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def add_para(doc, text="", *, bold_prefix=None, style=None, italic=False, color=DARK):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic, color=color)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Number")


def add_callout(doc, title: str, text: str, fill=LIGHT_BLUE, title_color=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(text), color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9.5, bold=True, color=DARK)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(text)), size=9.5, color=DARK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_requirement(doc, req_id: str, title: str, priority: str, requirement: str, criteria: list[str]) -> None:
    p = doc.add_paragraph(style="Heading 3")
    set_run(p.add_run(f"{req_id} - {title}"), size=12, bold=True, color=BLUE)
    tag = p.add_run(f"  [{priority}]")
    set_run(tag, size=9, bold=True, color="7A5A00" if priority == "Should" else "9B1C1C")
    add_para(doc, f"Yêu cầu: {requirement}", bold_prefix="Yêu cầu: ")
    for criterion in criteria:
        add_para(doc, criterion, style="List Bullet")


def add_section_break(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("NXT Automated Trading System | Business Requirements"), size=9, bold=True, color=MID_GRAY)
    add_page_number(section.footer.paragraphs[0])

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    title = doc.add_paragraph(style="Title")
    set_run(title.add_run("BUSINESS REQUIREMENTS DOCUMENT"), size=24, bold=True, color=DARK)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_run(subtitle.add_run("Hệ thống tự động scan tín hiệu, đề xuất lệnh, phê duyệt và giám sát giao dịch theo NXT latest"), size=13, color=MID_GRAY)

    metadata = [
        ["Thuộc tính", "Giá trị"],
        ["Phiên bản", "v0.2 - USD-M rulebook and desktop-first baseline"],
        ["Ngày", date.today().strftime("%d/%m/%Y")],
        ["Business owner", "Người dùng/Chủ tài khoản giao dịch - cần xác nhận"],
        ["Phạm vi baseline", "NXT v3.5 USD-M; BTCUSDT, BNBUSDT, SOLUSDT; USD-M perpetual 1D; promoted rulebook 2026-07-14"],
        ["Trạng thái", "Draft for review - chưa phê duyệt live trading"],
    ]
    add_table(doc, metadata[0], metadata[1:], [1800, 7560], LIGHT_BLUE)

    add_callout(
        doc,
        "Khuyến nghị kiểm soát chính",
        "MVP phải chạy Binance testnet/paper mode trước. Mọi lệnh mở vị thế cần người dùng Approve/Reject; chỉ được bật live execution sau khi UAT, reconciliation và kill switch đạt yêu cầu.",
        fill=AMBER,
        title_color="7A5A00",
    )

    doc.add_heading("1. Tóm tắt điều hành", level=1)
    add_para(
        doc,
        "Mục tiêu là phát triển hệ thống giao dịch bán tự động dựa trên NXT latest, có khả năng scan tín hiệu hằng ngày, lưu lịch sử, thông báo, sinh đề xuất order, nhận quyết định của người dùng, gửi lệnh lên Binance và theo dõi vòng đời vị thế. Hệ thống phải ưu tiên kiểm soát rủi ro, khả năng truy vết và ngăn hành động trùng lặp hơn tốc độ thực thi."
    )
    add_bullets(doc, [
        "Current state đã có desktop scanner/app dùng chung NXT v3.5 USD-M, Telegram notification, lịch chạy Windows và lịch sử signal JSON.",
        "Khoảng trống chính: chưa có approval workflow, Binance authenticated trading, order/position state, event reconciliation, daily action recommendation và operational controls.",
        "Target state là human-in-the-loop: hệ thống đề xuất; người dùng quyết định; hệ thống chỉ thực thi đúng order plan đã được duyệt và tiếp tục monitor.",
        "Promoted latest sử dụng Binance USD-M perpetual 1D candles và USD-M historical funding; live execution vẫn chỉ mở sau testnet/UAT và owner approval.",
    ])

    doc.add_heading("2. Bối cảnh và current state", level=1)
    add_table(doc, ["Năng lực", "Hiện trạng", "Khoảng trống"], [
        ["NXT rule engine", "NXT v3.5 USD-M 1D; SSL14; Runner A; Early-BE 7%; G-01 anti-reversal; G-02 block SHORT sau losing pre-TP1 LONG SSL exit; LONG-only continuation.", "Cần version locking, source artifact hash, golden regression và migration khi promote latest."],
        ["Universe", "BTCUSDT, BNBUSDT, SOLUSDT.", "Cần cấu hình symbol và kiểm tra contract availability."],
        ["Scan", "Daily scan desktop qua Binance USD-M data API, lịch Windows 07:10 ICT.", "Cần SLA, retry, health alert, late-candle handling và stale-data block."],
        ["Notification", "Telegram entry alert/no-signal tùy cấu hình.", "Cần Approve/Reject, expiry, escalation và delivery status."],
        ["History", "Signal history JSON và last scan JSON.", "Cần database/audit log, lifecycle status, immutable events."],
        ["Order plan", "Đề xuất entry, initial stop, TP1, breakeven, runner exit.", "Chưa validate exchange filters, balance, leverage, partial fills."],
        ["Execution", "Chưa gửi authenticated trading order.", "Cần testnet/live gateway, idempotency, API security."],
        ["Monitoring", "Rule engine tái dựng trạng thái lý thuyết.", "Cần sync order/fill/position thật và recommendation hằng ngày."],
    ], [1800, 3600, 3960])

    doc.add_heading("3. Mục tiêu kinh doanh và KPI", level=1)
    add_table(doc, ["ID", "Mục tiêu/KPI", "Mức mục tiêu ban đầu"], [
        ["OBJ-01", "Không bỏ sót scan daily sau khi nến Binance 1D đóng.", ">= 99% ngày; retry và alert nếu quá 07:20 ICT."],
        ["OBJ-02", "Mọi signal hợp lệ được lưu và thông báo một lần.", "100% idempotent theo signal ID."],
        ["OBJ-03", "Người dùng có đủ dữ liệu để quyết định.", "Order plan, risk, price deviation, expiry, rationale."],
        ["OBJ-04", "Không có lệnh live khi chưa được phê duyệt.", "100% entry order có approval record hợp lệ."],
        ["OBJ-05", "Order và position nội bộ khớp Binance.", "Reconciliation không có mismatch chưa xử lý quá 15 phút."],
        ["OBJ-06", "Mọi thay đổi SL/TP/close đều truy vết được.", "100% action có actor, timestamp, source và Binance IDs."],
    ], [1100, 5060, 3200])

    doc.add_heading("4. Stakeholders và vai trò", level=1)
    add_table(doc, ["Vai trò", "Trách nhiệm", "Quyền chính"], [
        ["Trader/Owner", "Xem tín hiệu, quyết định vào lệnh, override/close, chấp nhận rủi ro.", "Approve, Reject, Cancel, Emergency Close."],
        ["NXT Rule Owner", "Quản lý version rule, xác nhận promotion và regression.", "Publish/retire rule version."],
        ["System Operator", "Theo dõi scheduler, integration, lỗi và reconciliation.", "Retry, pause, investigate; không tự ý approve trade."],
        ["System", "Scan, validate, đề xuất, thực thi đã duyệt, monitor và audit.", "Chỉ hành động trong policy đã cấu hình."],
        ["Binance", "Nguồn market/account data và execution venue.", "Xác thực, matching, order/fill/position status."],
    ], [1800, 4700, 2860])

    doc.add_heading("5. Phạm vi", level=1)
    doc.add_heading("5.1 In scope", level=2)
    add_bullets(doc, [
        "Daily scan theo NXT latest version đã publish; lưu indicator snapshot và lý do pass/fail.",
        "Signal history, deduplication và trạng thái từ Detected đến Closed/Expired/Rejected.",
        "Telegram và local app cho notification, review, Approve/Reject và status.",
        "Order recommendation gồm side, entry, quantity, initial SL, TP1, runner rule, expected risk và expiry.",
        "Binance USD-M Futures testnet integration; live integration chỉ mở qua feature flag và approval gate.",
        "Order placement, cancellation, fill tracking, partial fill handling, protective order placement và reconciliation.",
        "Daily position monitoring và event-driven updates khi order/fill thay đổi.",
        "Khuyến nghị TP1, move SL to breakeven, runner exit, emergency action và exception handling.",
        "Audit log, health monitoring, backup/export và operational controls.",
        "Vận hành desktop-first trên Windows: local app và Windows Task Scheduler là runtime MVP chính thức.",
    ])
    doc.add_heading("5.2 Out of scope cho MVP", level=2)
    add_bullets(doc, [
        "Tự động thay đổi NXT rules bằng AI hoặc tự promote strategy.",
        "High-frequency/intraday strategy, copy trading, multi-exchange routing.",
        "Quản lý tài sản ngoài BTC/BNB/SOL nếu chưa cấu hình và UAT.",
        "Tự động rút tiền, chuyển tiền hoặc cấp quyền withdrawal.",
        "Cam kết lợi nhuận hoặc coi backtest là dự báo kết quả tương lai.",
        "Auto-approve entry live hoàn toàn không có người dùng trong vòng kiểm soát MVP.",
        "Cloud deployment, cloud scheduler hoặc remote execution service trong MVP; chỉ đánh giá sau desktop testnet/paper pilot ổn định.",
    ])
    doc.add_heading("5.3 Chưa quyết định", level=2)
    add_bullets(doc, [
        "Thiết lập tài khoản USD-M Futures (testnet/live), jurisdiction và contract availability cho từng symbol.",
        "Risk sizing live: fixed USD per R hay % equity; mức default và hard cap.",
        "Leverage, margin mode, price deviation tolerance và approval expiry.",
        "Có cho phép approve qua Telegram inline button hay chỉ qua app có xác thực.",
        "Mức tự động hóa sau MVP: auto-protective orders, auto-TP/SL, auto-runner exit.",
        "Cloud-readiness criteria sau desktop pilot; đây không phải quyết định blocking cho MVP.",
    ])

    doc.add_heading("6. Quy trình nghiệp vụ target state", level=1)
    doc.add_heading("6.1 Luồng scan và phê duyệt entry", level=2)
    add_numbered(doc, [
        "Windows Task Scheduler khởi chạy desktop scanner sau khi Binance USD-M 1D candle đóng (mặc định 07:10 ICT); hệ thống kiểm tra candle completeness, data source, contract và rule version.",
        "Rule engine scan BTCUSDT, BNBUSDT, SOLUSDT và sinh signal snapshot nếu đạt NXT latest.",
        "Risk engine lấy account snapshot, open exposure, exchange filters và policy để tính order proposal.",
        "Hệ thống lưu signal/order proposal trước khi gửi notification; mỗi proposal có ID duy nhất và thời hạn.",
        "Người dùng chọn Approve hoặc Reject. Approve phải hiển thị lại symbol, side, quantity, risk, SL, TP1 và môi trường testnet/live.",
        "Execution gateway revalidate giá, balance/margin, position, filters và approval trước khi gửi lệnh.",
        "Khi entry fill, hệ thống tạo protective orders, cập nhật position và thông báo kết quả.",
        "Nếu bất kỳ bước nào không chắc chắn, hệ thống chuyển trạng thái Needs Review thay vì đoán hoặc gửi lại lệnh.",
    ])

    doc.add_heading("6.2 Luồng monitor vị thế", level=2)
    add_numbered(doc, [
        "Nhận real-time order/account events từ Binance và chạy reconciliation định kỳ bằng REST.",
        "Đánh giá trạng thái: entry pending, open, TP1 hit, runner active, stopped, closed hoặc mismatch.",
        "Sau mỗi daily close, rule engine tính khuyến nghị tiếp theo dựa trên position thật và NXT latest.",
        "Nếu TP1 fill, hệ thống xác nhận remaining quantity rồi đề xuất/đặt stop tại breakeven theo policy.",
        "Nếu opposite SSL flip, hệ thống đề xuất close runner; chỉ tự close nếu policy đã được phê duyệt cho loại action này.",
        "Mọi action, rejection, Binance response và manual override được ghi audit event.",
    ])

    doc.add_heading("6.3 Signal lifecycle", level=2)
    add_table(doc, ["Trạng thái", "Ý nghĩa", "Chuyển tiếp hợp lệ"], [
        ["Detected", "Rule đạt, đã lưu.", "Proposed, Invalidated"],
        ["Proposed", "Đã có order plan.", "Pending Approval, Invalidated"],
        ["Pending Approval", "Đã thông báo, chờ user.", "Approved, Rejected, Expired"],
        ["Approved", "Approval còn hiệu lực.", "Submitting, Cancelled"],
        ["Submitting", "Đang gửi Binance.", "Entry Pending, Needs Review"],
        ["Entry Pending", "Order accepted/chưa fill đủ.", "Open, Cancelled, Needs Review"],
        ["Open", "Có position thật.", "TP1 Hit, Stopped, Closing, Needs Review"],
        ["TP1 Hit", "Đóng 50%; runner còn lại.", "Runner Active, Needs Review"],
        ["Runner Active", "SL tại BE, chờ SSL flip/stop.", "Closing, Closed, Stopped"],
        ["Closed", "Position kết thúc và reconciled.", "Final"],
        ["Rejected/Expired", "Không được thực thi.", "Final hoặc Re-propose mới"],
        ["Needs Review", "Trạng thái không chắc chắn/mismatch.", "Manual resolution"],
    ], [1800, 3600, 3960])

    doc.add_heading("7. Business rules", level=1)
    rules = [
        ("BR-01", "Rule version", "Mỗi signal phải gắn immutable NXT version, parameter snapshot, data-source/contract và source artifact hash; thay latest không được thay đổi signal lịch sử."),
        ("BR-02", "Universe", "Mặc định BTCUSDT, BNBUSDT, SOLUSDT; symbol ngoài danh sách phải được cấu hình, backtest và phê duyệt."),
        ("BR-03", "Signal timing and source", "Chỉ dùng Binance USD-M perpetual 1D candle đã đóng tại 00:00 UTC; không phát signal từ candle đang chạy hoặc cache stale."),
        ("BR-04", "Primary LONG", "SSL14 bullish flip; EMA20 cross-up trong candle signal hoặc 2 candles trước; distance EMA50 <= 2 ATR14; RSI14 > 50."),
        ("BR-05", "Primary SHORT", "SSL14 bearish flip; EMA20 cross-down trong candle signal hoặc 2 candles trước; distance EMA50 <= 2 ATR14; RSI14 < 50."),
        ("BR-06", "LONG Continuation", "Chỉ LONG continuation khi SSL14 bullish flip, close > EMA20 > EMA50, có Low <= EMA20 trong 5 candles gần nhất, close > EMA20 và close > close candle trước. SHORT continuation tắt."),
        ("BR-07", "Profitable SSL reversal guard (G-01)", "Sau SSL runner exit có net R trước funding >= +0.50R, block entry ngược chiều tại exit candle và candle kế tiếp; sau LONG block SHORT Primary, sau SHORT block LONG Primary/Continuation."),
        ("BR-08", "Initial risk and TP1", "Entry tại open candle kế tiếp; initial stop = 1.5 x ATR14 signal. TP1 = 2.5 x ATR14 signal; đóng 50% tại TP1 rồi stop phần còn lại về entry."),
        ("BR-09", "Early-BE and exit order", "Từ candle đầu tiên sau entry và trước TP1: LONG High >= Entry x 1.07 hoặc SHORT Low <= Entry x 0.93 kích hoạt stop về entry cho candle sau. Trên mỗi candle, kiểm tra stop trước, rồi TP1/Early-BE, rồi opposite SSL exit."),
        ("BR-10", "Losing pre-TP1 LONG guard (G-02)", "Khi prior LONG chưa TP1 thoát đúng bằng SSL bearish flip với net R trước funding < 0, block chỉ SHORT Primary tại exit candle và candle kế tiếp; không block LONG setup."),
        ("BR-11", "Indicator calculation", "EMA20/EMA50 tính trên close; ATR14 = SMA(True Range,14); RSI14 smoothing Wilder; SSL14 dùng SMA14 High/Low và giữ state khi close nằm giữa bands."),
        ("BR-12", "Approval", "Không gửi entry order live nếu không có approval còn hiệu lực, đúng proposal hash và đúng environment."),
        ("BR-13", "Risk cap", "Order bị chặn nếu vượt per-trade, symbol, portfolio, leverage hoặc daily-loss limits."),
        ("BR-14", "Idempotency", "Một approved proposal chỉ được tạo tối đa một logical entry order; retry phải dùng client order ID ổn định."),
        ("BR-15", "Protective orders", "Sau entry fill, SL phải được tạo ngay; nếu không tạo được, position chuyển Critical và kích hoạt policy xử lý khẩn cấp."),
        ("BR-16", "Reconciliation", "Binance là source of truth cho order/fill/position; khác biệt phải được ghi nhận và khóa automation liên quan."),
        ("BR-17", "Manual override", "Manual close/cancel trên Binance hoặc app phải được phát hiện; hệ thống không được tự tái mở vị thế vì cho rằng lệnh thiếu."),
    ]
    add_table(doc, ["ID", "Tên", "Quy tắc"], [[a, b, c] for a, b, c in rules], [1100, 1900, 6360])

    doc.add_heading("8. Functional requirements", level=1)
    doc.add_heading("8.1 Requirement catalogue", level=2)
    add_requirement(doc, "FR-01", "Scheduled scanning", "Must",
                    "Hệ thống shall tự động chạy daily scan sau khi nến 1D đóng và cho phép scan thủ công.",
                    ["Given scheduler active, when đến giờ cấu hình, then scan chạy một lần với run ID.",
                     "Nếu data chưa complete, hệ thống retry theo policy và không dùng candle chưa đóng.",
                     "Kết quả phải lưu checked symbols, candle date, USD-M contract/data source, rule version, artifact hash, errors và duration."])
    add_requirement(doc, "FR-02", "Signal detection and persistence", "Must",
                    "Hệ thống shall lưu signal hợp lệ trước notification và không tạo duplicate.",
                    ["Signal ID phải deterministic theo symbol, signal date, side, type, rule version và source artifact hash.",
                     "Signal snapshot phải chứa indicator values, data source/contract, entry reference, SL, TP1, guard decision pass/fail và detected timestamp.",
                     "Re-run cùng dữ liệu không tạo thêm signal record."])
    add_requirement(doc, "FR-03", "Order proposal", "Must",
                    "Hệ thống shall tạo order proposal dựa trên signal, account state, risk policy và Binance filters.",
                    ["Quantity được làm tròn đúng step size; price đúng tick size; notional đạt filter.",
                     "Proposal hiển thị estimated risk, notional, leverage, fees/funding caveat và price deviation.",
                     "Proposal không hợp lệ phải ghi lý do Blocked, không cho Approve."])
    add_requirement(doc, "FR-04", "User decision", "Must",
                    "Hệ thống shall cho người dùng Approve hoặc Reject proposal và ghi bằng chứng quyết định.",
                    ["Approve yêu cầu explicit confirmation và hiển thị testnet/live.",
                     "Approval hết hạn khi quá cutoff, proposal thay đổi hoặc giá lệch quá tolerance.",
                     "Reject yêu cầu optional reason và không được gửi lệnh."])
    add_requirement(doc, "FR-05", "Binance pre-trade validation", "Must",
                    "Hệ thống shall revalidate ngay trước execution.",
                    ["Kiểm tra API permission, server time, symbol status, account/margin, open position, filters và risk limits.",
                     "Kiểm tra proposal hash và chưa có order cùng idempotency key.",
                     "Bất kỳ check thất bại nào chuyển Needs Review/Blocked."])
    add_requirement(doc, "FR-06", "Entry execution", "Must",
                    "Hệ thống shall gửi entry order đã duyệt và lưu đầy đủ request/response đã mask secret.",
                    ["Testnet phải hỗ trợ validate/test order trước live.",
                     "Network timeout không được tự suy luận lệnh thất bại; phải query bằng client order ID.",
                     "Partial fill phải cập nhật actual quantity và protective orders tương ứng."])
    add_requirement(doc, "FR-07", "Protective orders", "Must",
                    "Hệ thống shall quản lý initial SL và TP1 theo actual filled quantity.",
                    ["SL/TP phải reduce-only khi execution product hỗ trợ.",
                     "Sau TP1 fill, cancel/replace stop để bảo vệ remaining quantity tại breakeven.",
                     "Nếu protective order reject, gửi Critical alert và thực thi fail-safe policy."])
    add_requirement(doc, "FR-08", "Position monitoring", "Must",
                    "Hệ thống shall đồng bộ order, fill và position real-time kết hợp polling/reconciliation.",
                    ["User data disconnect phải reconnect và backfill missed events.",
                     "Daily monitor phải tạo action recommendation kể cả không có action.",
                     "Manual Binance activity phải được nhận diện và reflected trong internal state."])
    add_requirement(doc, "FR-09", "Exit recommendations", "Must",
                    "Hệ thống shall đề xuất hoặc thực hiện action đúng NXT policy và automation level.",
                    ["TP1, Early-BE, BE move, SSL runner exit, stop/close phải có rationale, rule version và current position snapshot.",
                     "Position mở giữ entry rule version; latest rule mới chỉ áp dụng khi có explicit migration record hoặc policy đã được phê duyệt.",
                     "Action cần approval phải không được gửi trước approval.",
                     "Không được close quantity vượt position thực."])
    add_requirement(doc, "FR-10", "Notifications", "Must",
                    "Hệ thống shall gửi Telegram cho signal, approval status, fills, protective-order failures và daily recommendations.",
                    ["Notification có correlation ID và link/command tới proposal.",
                     "Delivery failure được retry và ghi status.",
                     "Critical alerts được gửi lặp/escalate đến khi acknowledged hoặc resolved."])
    add_requirement(doc, "FR-11", "Dashboard and history", "Should",
                    "Local app shall hiển thị signal, proposal, decision, order, position, action và audit timeline.",
                    ["Có filter theo symbol/status/date/environment.",
                     "Hiển thị rõ data source và last successful sync.",
                     "Cho phép export CSV/JSON nhưng không lộ secrets."])
    add_requirement(doc, "FR-12", "Operational controls", "Must",
                    "Hệ thống shall có pause, kill switch và environment gate.",
                    ["Pause ngăn entry mới nhưng vẫn monitor/protect position.",
                     "Kill switch ngăn mọi automated action trừ policy đóng khẩn cấp được owner xác nhận.",
                     "Live mode yêu cầu explicit config, valid credentials và deployment approval."])

    doc.add_heading("9. Non-functional requirements", level=1)
    add_table(doc, ["ID", "Nhóm", "Yêu cầu"], [
        ["NFR-01", "Security", "API key không có withdrawal; tách read/trade key nếu khả thi; IP allowlist; secrets không lưu/log plaintext."],
        ["NFR-02", "Reliability", "Retry có exponential backoff; idempotent execution; recover sau restart mà không mất state."],
        ["NFR-03", "Availability", "Desktop-first: daily scan target >=99% trên Windows Task Scheduler; monitor local service health và alert khi scheduler missed."],
        ["NFR-04", "Performance", "Scan hoàn tất trong 5 phút; approval action phản hồi UI <3 giây; execution request theo exchange SLA."],
        ["NFR-05", "Auditability", "Append-only business event log; timestamp UTC; actor, correlation ID, rule version, before/after state."],
        ["NFR-06", "Data integrity", "Atomic writes/transactions; backup; checksum/version; reconciliation với Binance."],
        ["NFR-07", "Observability", "Structured logs, metrics, health endpoint, error taxonomy; không ghi token/signature."],
        ["NFR-08", "Maintainability", "Signal engine dùng chung giữa backtest/scanner/app; automated regression theo golden trades."],
        ["NFR-09", "Usability", "Màn hình/Telegram phân biệt rõ Testnet và Live; action nguy hiểm cần confirmation."],
        ["NFR-10", "Time", "Lưu UTC; hiển thị Asia/Saigon; sync Binance server time cho signed requests."],
    ], [1100, 1700, 6560])

    doc.add_heading("10. Data requirements", level=1)
    add_table(doc, ["Entity", "Trường chính", "Retention/Control"], [
        ["RuleVersion", "version, parameters, USD-M contract/session, source artifact hash, promotedAt", "Immutable; retain indefinitely."],
        ["ScanRun", "runId, start/end, candleDate, symbols, status, errors", ">= 2 years."],
        ["Signal", "signalId, ruleVersion, indicators, rationale, detectedAt", "Immutable snapshot."],
        ["OrderProposal", "proposalId, signalId, quantity, levels, risk, expiry, hash", "Versioned; no in-place mutation after approval."],
        ["Decision", "approve/reject, actor, timestamp, channel, proposalHash", "Append-only."],
        ["ExchangeOrder", "clientOrderId, Binance orderId, request status, fills", "Source-linked; mask credentials."],
        ["Position", "symbol, side, actual qty, entry, realized/unrealized, lifecycle", "Reconciled snapshot + events."],
        ["Recommendation", "type, rationale, requiredBy, automationLevel, status", "Append-only with outcome."],
        ["AuditEvent", "actor, eventType, object, before/after, correlationId", "Tamper-evident; retain indefinitely."],
    ], [1700, 4860, 2800])

    doc.add_heading("11. Binance integration requirements", level=1)
    add_bullets(doc, [
        "Technical baseline: Binance USD-M Futures because promoted NXT latest uses USD-M contract candles, supports SHORT and includes funding. Owner must still approve account/product eligibility before live execution.",
        "SIGNED requests use API key/signature and server-time-safe timestamp/recvWindow handling.",
        "Use unique client order IDs for idempotency and recovery after ambiguous network errors.",
        "Read exchange information before sizing; enforce price, quantity, notional and order-count filters.",
        "Use user data stream for account/order events and REST reconciliation after reconnect.",
        "Handle rate-limit headers, HTTP 429 backoff and 418 protection; avoid aggressive polling.",
        "Support test order/testnet before live; environment endpoints and credentials must never be mixed.",
        "Store Binance order IDs, client IDs, fills, commissions, realized P&L and funding where available.",
    ])
    add_callout(
        doc,
        "Fail-safe principle",
        "Khi response không chắc chắn (timeout/5xx/disconnect), hệ thống phải query trạng thái order bằng client order ID trước khi retry. Không được gửi một lệnh mới chỉ vì chưa nhận được response.",
        fill=RED,
        title_color="9B1C1C",
    )

    doc.add_heading("12. Risk management và controls", level=1)
    add_table(doc, ["Control", "Khuyến nghị mặc định", "Cần phê duyệt"], [
        ["Environment", "Paper/Testnet trước; Live disabled.", "Điều kiện go-live và owner."],
        ["Per-trade risk", "1% equity default; hard cap 2%.", "Không tự dùng backtest 5%/trade."],
        ["Portfolio risk", "Giới hạn tổng open risk và correlated exposure.", "Mức % cụ thể."],
        ["Margin mode", "Isolated.", "Xác nhận product/account mode."],
        ["Leverage", "Thấp và configurable; không dùng để tăng risk budget.", "Mức tối đa theo symbol."],
        ["Daily loss limit", "Pause entry mới khi chạm limit.", "Mức R/% và reset time."],
        ["Price deviation", "Re-approve nếu giá khác proposal quá tolerance.", "Tolerance theo % hoặc ATR."],
        ["Approval expiry", "Hết hạn sau cutoff gần daily open.", "Thời lượng cụ thể."],
        ["Protective failure", "Critical alert; block entry mới; manual/fail-safe close.", "Policy tự close hay chờ owner."],
        ["Kill switch", "Một thao tác, audit đầy đủ.", "Ai có quyền kích hoạt/khôi phục."],
    ], [2200, 3960, 3200])

    doc.add_heading("13. Reporting và notifications", level=1)
    add_bullets(doc, [
        "Daily scan summary: run status, symbols checked, new signals, errors, data freshness.",
        "Signal proposal: rule/version, signal type, side, levels, size, risk, expiry, Approve/Reject.",
        "Execution report: accepted/rejected, fill quantity/price, slippage, Binance IDs, protective order state.",
        "Position daily report: current qty, entry, SL, TP1, P&L, funding, NXT state, recommended action.",
        "Exception report: mismatch, disconnected stream, stale data, rejected protective order, rate limit.",
        "Weekly audit: signals detected vs decisions vs executions vs outcomes; manual overrides; missed jobs.",
    ])

    doc.add_heading("14. Error và exception scenarios", level=1)
    add_table(doc, ["Tình huống", "Hành vi yêu cầu"], [
        ["Binance market data unavailable", "Retry; dùng cache chỉ để hiển thị, không phát signal mới từ stale data; alert operator."],
        ["Telegram unavailable", "Signal vẫn lưu; retry delivery; app vẫn cho review."],
        ["Approval nhận hai lần", "Idempotent; chỉ một Decision hợp lệ và một logical order."],
        ["Price gap qua SL trước entry", "Invalidate/recalculate theo policy; không gửi order cũ."],
        ["Entry partial fill", "Protect filled qty; theo dõi remaining; proposal/order state phản ánh partial."],
        ["SL/TP reject", "Critical; khóa automation entry mới; reconcile và áp dụng fail-safe."],
        ["User đóng thủ công trên Binance", "Phát hiện qua account event/reconciliation; cancel orphan protective orders."],
        ["User sửa order trên Binance", "Mark External Change; không overwrite tự động khi chưa review."],
        ["Service restart", "Khôi phục từ durable state; reconcile trước khi action."],
        ["Rule latest thay đổi khi có position", "Position giữ entry rule version; exit policy phải theo decision đã phê duyệt hoặc migration record."],
    ], [2600, 6760])

    doc.add_heading("15. Acceptance và UAT", level=1)
    add_bullets(doc, [
        "UAT-01 Golden signal: scanner khớp danh sách signal NXT latest trên tập USD-M 1D kiểm soát, theo exact rule version và artifact hash.",
        "UAT-02 G-02 regression: prior LONG chưa TP1 + SSL bearish exit + net R trước funding < 0 phải block SHORT Primary tại exit candle và candle kế tiếp; không block ở candle thứ ba, khi TP1 đã hit, khi exit không phải SSL bearish, hoặc khi net R >= 0.",
        "UAT-03 Duplicate: chạy cùng candle nhiều lần không tạo signal/order trùng.",
        "UAT-04 Approval: reject/expired proposal không bao giờ gọi trading endpoint.",
        "UAT-05 Testnet E2E: signal -> proposal -> approve -> entry -> SL/TP -> fill -> close -> audit đầy đủ.",
        "UAT-06 Ambiguous response: timeout sau submit không tạo duplicate; hệ thống query order trước retry.",
        "UAT-07 Partial fill: protective quantity khớp actual fill và không vượt position.",
        "UAT-08 Manual intervention: đóng/sửa order ngoài hệ thống được phát hiện và reconciled.",
        "UAT-09 Kill switch: block entry mới nhưng vẫn cho read/monitor và hành động khẩn cấp theo policy.",
        "UAT-10 Recovery: restart giữa lifecycle không mất decision/order state.",
        "UAT-11 Security: logs/export không chứa API secret/signature; key không có withdrawal permission.",
        "UAT-12 Operational run: tối thiểu 30 ngày desktop paper/testnet không có Sev-1/duplicate order trước live review.",
    ])

    doc.add_heading("16. Roadmap đề xuất", level=1)
    add_table(doc, ["Giai đoạn", "Phạm vi", "Exit criteria"], [
        ["Phase 0 - Discovery", "Chốt open questions, product, risk policy, approval model.", "BRD sign-off và decision log."],
        ["Phase 1 - Desktop paper workflow", "Windows desktop local app, Task Scheduler, database/lifecycle, app/Telegram approval, simulated fills.", "UAT signal/history/approval đạt."],
        ["Phase 2 - Desktop USD-M testnet", "Authenticated gateway, orders, stream, reconciliation và controls trên desktop runtime.", "30 ngày ổn định; no duplicate."],
        ["Phase 3 - Desktop live assisted", "Live entry sau manual approval; auto protective orders trên desktop runtime.", "Go-live checklist và low-risk pilot."],
        ["Phase 4 - Managed automation", "Tùy chọn auto TP/SL/runner theo policy; chỉ đánh giá cloud-readiness sau desktop pilot.", "Risk committee/owner approval."],
    ], [1900, 4500, 2960])

    doc.add_heading("17. Dependencies và risks", level=1)
    add_table(doc, ["Risk/Dependency", "Ảnh hưởng", "Mitigation"], [
        ["NXT latest thay đổi", "Signal live lệch backtest.", "Version locking, golden regression, controlled promotion."],
        ["Windows desktop tắt/mất mạng", "Missed scan/monitor.", "StartWhenAvailable; health alert; runbook restart/reconcile. Cloud chỉ đánh giá sau desktop pilot ổn định."],
        ["Binance API/region/account limits", "Không thể trade hoặc stream.", "Test account capability; endpoint abstraction; retry/reconcile."],
        ["Risk sizing quá cao", "Drawdown/liquidation.", "Configurable % equity, hard caps, isolated margin."],
        ["Duplicate/ambiguous orders", "Overexposure.", "Idempotency, client IDs, query-before-retry."],
        ["Secret compromise", "Unauthorized trading.", "Least privilege, IP allowlist, rotation, secret store."],
        ["Manual exchange actions", "Internal state sai.", "User data stream + periodic reconciliation."],
        ["Backtest/live divergence", "Kỳ vọng sai.", "Slippage/funding logging, live performance attribution."],
    ], [2600, 2860, 3900])

    doc.add_heading("18. Open questions cần business owner trả lời", level=1)
    questions = [
        ["OQ-01", "Tài khoản USD-M Futures có đủ eligibility/contract availability cho testnet và live không?", "Owner", "Blocking"],
        ["OQ-02", "Risk live mỗi trade là fixed USD hay % equity? Default và hard cap?", "Owner", "Blocking"],
        ["OQ-03", "Leverage tối đa và margin mode cho từng symbol?", "Owner", "Blocking"],
        ["OQ-04", "Entry cần approve qua Telegram, app hay cả hai? Có yêu cầu PIN/2FA?", "Owner/Security", "Blocking"],
        ["OQ-05", "Proposal hết hạn sau bao lâu và price deviation tolerance là bao nhiêu?", "Owner", "Blocking"],
        ["OQ-06", "Sau entry fill, SL/TP có được auto-place không cần approve lần hai?", "Owner", "Blocking"],
        ["OQ-07", "Runner exit do SSL flip là recommendation hay auto-close?", "Owner", "High"],
        ["OQ-08", "Policy khi không đặt được stop: auto-close market hay chờ manual?", "Owner", "Blocking"],
        ["OQ-09", "Daily loss/open-risk limits và correlated exposure cap?", "Owner", "Blocking"],
        ["OQ-10", "Ai có quyền bật Live, pause, kill switch và resume?", "Owner", "High"],
        ["OQ-11", "Retention/audit backup cần bao lâu và lưu ở đâu?", "Owner/Operator", "Medium"],
        ["OQ-12", "Sau desktop testnet/paper pilot, tiêu chí nào kích hoạt đánh giá cloud-readiness?", "Owner", "Deferred"],
    ]
    add_table(doc, ["ID", "Câu hỏi", "Owner", "Mức"], questions, [1000, 5800, 1500, 1060])

    doc.add_heading("19. Decision log ban đầu", level=1)
    add_table(doc, ["Ngày", "Decision/Assumption", "Rationale", "Status"], [
        [date.today().isoformat(), "Baseline là NXT v3.5 USD-M BTC/BNB/SOL 1D; promoted rulebook 2026-07-14.", "Khớp latest summary, rulebook và signal-level regression đã publish.", "Confirmed current state"],
        [date.today().isoformat(), "Human approval cho entry trong MVP.", "Giảm rủi ro giao dịch live ngoài ý muốn.", "Recommended"],
        [date.today().isoformat(), "Testnet/paper trước live.", "Cần chứng minh execution và reconciliation.", "Recommended"],
        [date.today().isoformat(), "USD-M Futures là technical baseline cho data/backtest và execution integration.", "Promoted latest dùng USD-M candles/funding; live eligibility vẫn cần owner approval.", "Confirmed technical baseline"],
        [date.today().isoformat(), "Live risk không kế thừa mặc định 1R=$1,000.", "Backtest account tương đương 5%/trade, rủi ro cao.", "Pending owner"],
        [date.today().isoformat(), "Desktop-first runtime cho MVP.", "Local app + Windows Task Scheduler là runtime chính thức; cloud không thuộc MVP.", "Confirmed"],
    ], [1400, 3400, 3360, 1200])

    doc.add_heading("20. Traceability summary", level=1)
    add_table(doc, ["Business objective", "Requirements", "UAT"], [
        ["OBJ-01/02 Scan và không trùng", "FR-01, FR-02, NFR-02", "UAT 1, 2, 9"],
        ["OBJ-03 Quyết định có đủ dữ liệu", "FR-03, FR-04, FR-10, FR-11", "UAT 3, 4"],
        ["OBJ-04 Không trade khi chưa duyệt", "BR-12, FR-04, FR-05, FR-12", "UAT 3, 8"],
        ["OBJ-05 Khớp Binance", "BR-14-16, FR-06-09", "UAT 4-7, 9"],
        ["OBJ-06 Audit đầy đủ", "FR-10-12, NFR-05-07", "UAT 4, 10"],
    ], [2300, 4100, 2960])

    doc.add_heading("Phụ lục A - Nguồn tham chiếu", level=1)
    add_bullets(doc, [
        "Workspace: latest/NXT_Latest_Summary.md - NXT v3.5 USD-M portfolio baseline.",
        "Workspace: latest/NXT_Latest_NXT35_USDM_BlockShortAfterLosingLong_System_And_Indicators.docx - promoted detailed rulebook.",
        "Workspace: latest/NXT_Latest_NXT35_USDM_BlockShortAfterLosingLong_SignalRegression.json - signal-level regression evidence for G-02.",
        "Workspace: app/nxt_signal_app.py - shared scan core, signal history và suggested orders.",
        "Workspace: scripts/daily_nxt_signal_scan.py - Telegram scheduled scanner.",
        "Binance Developer Docs - USD-M Futures introduction: https://developers.binance.com/en/docs/products/derivatives-trading-usds-futures/Introduction",
        "Binance Developer Docs - USD-M Futures common definitions and exchange information: https://developers.binance.com/en/docs/products/derivatives-trading-usds-futures/common-definition",
        "Binance Developer Docs - product catalog: https://developers.binance.com/en/docs/catalog",
    ])

    doc.add_heading("Phụ lục B - Definition of Done cho BRD", level=1)
    add_bullets(doc, [
        "Business owner trả lời toàn bộ câu hỏi Blocking.",
        "Decision log được cập nhật với product, risk, approval và fail-safe policy.",
        "Functional requirements được Product/Engineering/QA review và estimate.",
        "Security review chấp nhận key permissions và secret handling.",
        "UAT plan có test data, testnet account, owners và evidence format.",
        "BRD được sign-off trước khi bật bất kỳ live trading permission nào.",
    ])

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        if not section.header.paragraphs[0].text:
            set_run(section.header.paragraphs[0].add_run("NXT Automated Trading System | Business Requirements"), size=9, bold=True, color=MID_GRAY)
        if not section.footer.paragraphs[0].text:
            add_page_number(section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "BRD - NXT Automated Signal and Trading System"
    props.subject = "Business requirements for NXT signal scanning, approval, Binance execution and position monitoring"
    props.author = "Business Analysis"
    props.keywords = "NXT, Binance, BRD, trading automation, signal, approval, monitoring"
    props.comments = "Draft v0.2 for stakeholder review. Desktop-first MVP; not authorization for live trading."

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
