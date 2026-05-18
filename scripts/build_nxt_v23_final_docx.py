from __future__ import annotations

import json
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Workspace\Codex\Crypto trading ")
OUT_DIR = ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v23"
JSON_PATH = OUT_DIR / "nxt_crypto_btc_sol_sui_3y_v23_results.json"
XLSX_PATH = OUT_DIR / "NXT_Crypto_BTC_SOL_SUI_3Y_V23_Backtest.xlsx"
DOCX_PATH = OUT_DIR / "NXT_Trading_System_v2.3_Final.docx"


def read_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt_r(value: float) -> str:
    return f"{value:+.2f}R"


def fmt_pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def fmt_num(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill: str = "EAF0F6") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 41, 55)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    style_table(table)
    return table


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(17, 24, 39)


def summarize_trades(data: dict) -> dict:
    trades = sorted(data["trades"], key=lambda t: t["exitTime"])
    total = sum(t["rMultiple"] for t in trades)
    wins = sum(1 for t in trades if t["rMultiple"] > 0)
    eq = peak = 0.0
    max_dd = 0.0
    for t in trades:
        eq += t["rMultiple"]
        peak = max(peak, eq)
        max_dd = min(max_dd, eq - peak)
    return {
        "trades": len(trades),
        "wins": wins,
        "losses": len(trades) - wins,
        "winRate": wins / len(trades),
        "totalR": total,
        "avgR": total / len(trades),
        "bestR": max(t["rMultiple"] for t in trades),
        "worstR": min(t["rMultiple"] for t in trades),
        "maxDD": max_dd,
        "tp1Rate": sum(1 for t in trades if t.get("tp1Hit") == "Yes") / len(trades),
        "exitReasons": Counter(t["exitReason"] for t in trades),
        "long": [t for t in trades if t["side"] == "LONG"],
        "short": [t for t in trades if t["side"] == "SHORT"],
    }


def variant_total(path: Path) -> tuple[int, float, float, float] | None:
    if not path.exists():
        return None
    data = read_json(path)
    if "trades" not in data:
        return None
    s = summarize_trades(data)
    return s["trades"], s["totalR"], s["avgR"], s["maxDD"]


def build_doc() -> None:
    data = read_json(JSON_PATH)
    overall = summarize_trades(data)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12.5)

    header = section.header.paragraphs[0]
    header.text = "NXT Trading System v2.3"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("NXT Trading System v2.3").bold = True
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Bản chốt tạm sau backtest 3 năm trên BTC, SOL, SUI")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(75, 85, 99)

    created = datetime.now(timezone.utc).astimezone().strftime("%d-%b-%Y %H:%M")
    period_start = data["period"]["start"][:10]
    period_end = data["period"]["end"][:10]
    add_table(
        doc,
        ["Hạng mục", "Giá trị"],
        [
            ["Trạng thái", "Chốt tạm để theo dõi tiếp"],
            ["Nguồn rule", "NXT_Trading_System.docx + các điều chỉnh v2.3"],
            ["Dữ liệu", "Binance Spot daily klines"],
            ["Giai đoạn test", f"{period_start} đến {period_end}"],
            ["Danh mục", ", ".join(data["symbols"])],
            ["File chi tiết lệnh", str(XLSX_PATH)],
            ["Ngày xuất tài liệu", created],
        ],
        [4.0, 12.0],
    )

    add_heading(doc, "Kết luận chốt tạm", 1)
    for item in [
        "Giữ v2.3 làm bản chính vì tổng R và expectancy tốt nhất trong các nhánh đã test, trong khi drawdown vẫn nằm trong vùng chấp nhận được.",
        "Không dùng Weekly regime trong bản này, vì filter weekly làm giảm nhiều cơ hội và bỏ lỡ các nhịp đảo chiều daily có xác nhận volume.",
        "Các bản runner/trailing cải thiện một số lệnh bull-run riêng lẻ, nhưng tổng hệ thống 3 năm không vượt v2.3.",
        "Grid improve gần nhất giúp giảm drawdown, nhưng đánh đổi quá nhiều total R so với v2.3 gốc.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Thông số hệ thống v2.3", 1)
    add_table(
        doc,
        ["Nhóm rule", "Thiết lập chốt tạm"],
        [
            ["Timeframe", "Daily signal, entry tại open của nến daily kế tiếp."],
            ["Weekly regime", "OFF, không dùng làm filter entry."],
            ["Trend trigger", "SSL Channel 10/10 crossover. SSL xấp xỉ bằng SMA(high,10) và SMA(low,10)."],
            ["EMA confirmation", "Giá cross EMA20 trong vòng 3 nến gần nhất theo hướng lệnh."],
            ["EMA50 distance", "Khoảng cách từ close tới EMA50 <= 2 ATR(14)."],
            ["Volume filter", "Net Volume Binance: taker buy base volume x 2 - total volume. Long cần > 0, Short cần < 0."],
            ["Stop loss", "1.5 ATR(14) từ entry."],
            ["TP1", "1.5 ATR(14), chốt 50%, phần còn lại dời stop về breakeven."],
            ["TP2", "2.5 ATR(14), chốt 50% còn lại."],
            ["Exit khác", "Opposite SSL flip, stop loss, breakeven stop sau TP1, hoặc mark-to-market cuối kỳ."],
            ["Cost model", "0.06% fee + 0.05% slippage mỗi chiều, trừ trực tiếp vào R của từng lệnh."],
            ["Intraday assumption", "Nếu cùng ngày chạm nhiều mức, ưu tiên kiểm tra stop trước TP để bảo thủ."],
        ],
        [4.2, 11.8],
    )

    add_heading(doc, "Kết quả backtest 3 năm", 1)
    add_table(
        doc,
        ["Metric", "Giá trị"],
        [
            ["Tổng số lệnh", str(overall["trades"])],
            ["Win / Loss", f'{overall["wins"]} / {overall["losses"]}'],
            ["Win rate", fmt_pct(overall["winRate"])],
            ["Total R", fmt_r(overall["totalR"])],
            ["Average R / trade", fmt_r(overall["avgR"])],
            ["Best / Worst trade", f'{fmt_r(overall["bestR"])} / {fmt_r(overall["worstR"])}'],
            ["Max drawdown", fmt_r(overall["maxDD"])],
            ["Tỷ lệ hit TP1", fmt_pct(overall["tp1Rate"])],
        ],
        [5.0, 5.0],
    )

    add_heading(doc, "Đóng góp theo coin", 2)
    rows = []
    for s in data["summary"]:
        rows.append([
            s["symbol"],
            str(s["trades"]),
            f'{s["wins"]}/{s["losses"]}',
            fmt_pct(s["winRate"]),
            fmt_r(s["totalR"]),
            fmt_r(s["avgR"]),
            fmt_r(s["bestR"]),
            fmt_r(s["worstR"]),
        ])
    add_table(
        doc,
        ["Coin", "Trades", "W/L", "Win rate", "Total R", "Avg R", "Best", "Worst"],
        rows,
        [2.0, 1.5, 1.5, 2.0, 2.0, 2.0, 1.7, 1.7],
    )

    add_heading(doc, "Phân tích exit", 2)
    exit_rows = [[reason, str(count), fmt_pct(count / overall["trades"])] for reason, count in overall["exitReasons"].most_common()]
    add_table(doc, ["Exit reason", "Số lệnh", "Tỷ trọng"], exit_rows, [7.0, 2.0, 2.2])

    add_heading(doc, "So sánh với các nhánh đã thử", 1)
    variant_paths = [
        ("v2.2", ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v22" / "nxt_crypto_btc_sol_sui_3y_v22_results.json", "Bỏ weekly regime, rule cũ hơn v2.3."),
        ("v2.3", JSON_PATH, "Bản đang chốt tạm."),
        ("v2.4", ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v24" / "nxt_crypto_btc_sol_sui_3y_v24_results.json", "Runner sau TP1, trailing ATR."),
        ("v2.5", ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v25" / "nxt_crypto_btc_sol_sui_3y_v25_results.json", "Runner stop về TP1 rồi trail."),
        ("v2.6", ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v26" / "nxt_crypto_btc_sol_sui_3y_v26_results.json", "TP1 1.5 ATR, runner về BE rồi trail."),
        ("v3.0", ROOT / "outputs" / "nxt_crypto_btc_sol_sui_3y_v30" / "nxt_crypto_btc_sol_sui_3y_v30_results.json", "40/30/30 với runner."),
    ]
    comp_rows = []
    for name, path, note in variant_paths:
        v = variant_total(path)
        if v:
            trades, total_r, avg_r, max_dd = v
            comp_rows.append([name, str(trades), fmt_r(total_r), fmt_r(avg_r), fmt_r(max_dd), note])
    comp_rows.append(["Grid v2.3", "107", "+26.44R", "+0.25R", "-4.81R", "Distance <= 2.25 ATR, TP split 30/70; giảm DD nhưng total R thấp hơn v2.3."])
    add_table(doc, ["Version", "Trades", "Total R", "Avg R", "MaxDD", "Ghi chú"], comp_rows, [1.8, 1.5, 1.8, 1.7, 1.7, 7.5])

    add_heading(doc, "Case kiểm chứng 16-Oct-2023", 1)
    add_table(
        doc,
        ["Điểm kiểm tra", "Kết quả v2.3"],
        [
            ["Signal", "BTCUSDT LONG signal ngày 2023-10-16."],
            ["Entry", "Entry tại daily open ngày 2023-10-17, giá 28,500.77."],
            ["Lý do v2.3 bắt được lệnh", "EMA20 cross được chấp nhận trong vòng 3 nến gần nhất; distance tới EMA50 = 1.84 ATR, nằm trong ngưỡng <= 2 ATR."],
            ["Quản trị lệnh", "TP1 hit ngày 2023-10-20, TP2 hit ngày 2023-10-21."],
            ["Kết quả", "Ròng +1.28R sau fee/slippage."],
        ],
        [4.2, 11.8],
    )

    add_heading(doc, "Điểm cần theo dõi khi dùng thực chiến", 1)
    for item in [
        "Mẫu 3 coin vẫn chưa đủ để coi là ổn định trên toàn thị trường altcoin; cần forward test và mở rộng basket trước khi tăng risk.",
        "V2.3 phụ thuộc mạnh vào quality của Net Volume Binance; nếu data source thay đổi cần validate lại.",
        "Short side đóng góp R lớn hơn Long side trong mẫu này, nên không nên tự ý bỏ short khi chưa rerun.",
        "Các tín hiệu cùng ngày có thể chịu ảnh hưởng giả định thứ tự intraday; kết quả đã dùng cách bảo thủ nhưng vẫn cần kiểm chứng khi triển khai live.",
        "Risk per trade nên giữ cố định theo R cho đến khi có thêm mẫu forward test.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Tài liệu tham chiếu", 1)
    add_table(
        doc,
        ["Artifact", "Đường dẫn"],
        [
            ["Workbook chi tiết lệnh v2.3", str(XLSX_PATH)],
            ["JSON kết quả v2.3", str(JSON_PATH)],
            ["Script backtest v2.3", str(ROOT / "scripts" / "backtest_nxt_crypto_v23_3y.mjs")],
        ],
        [4.5, 11.5],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
