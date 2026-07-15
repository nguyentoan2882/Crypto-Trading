from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "HLD_NXT_Automated_Signal_Trading_System_v0.1.docx"
ASSET_DIR = ROOT / "outputs" / "hld_nxt_v01"
ARCHITECTURE_PNG = ASSET_DIR / "nxt_desktop_architecture.png"

NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "0F766E"
PALE_BLUE = "EAF2F8"
PALE_GREEN = "E8F3EC"
PALE_AMBER = "FFF4D6"
LIGHT_GRAY = "F2F4F7"
DARK = "172B4D"
MID = "667085"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\seguisb.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def image_color(value: str) -> str:
    return f"#{value}" if not value.startswith("#") else value


def draw_box(draw, rect, title: str, body: str, fill: str, outline: str = NAVY):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=16, fill=image_color(fill), outline=image_color(outline), width=2)
    draw.text((x1 + 16, y1 + 12), title, font=font(24, True), fill=image_color(NAVY))
    y = y1 + 48
    words = body.split()
    line = ""
    for word in words:
        next_line = f"{line} {word}".strip()
        if draw.textlength(next_line, font=font(18)) > (x2 - x1 - 32):
            draw.text((x1 + 16, y), line, font=font(18), fill=image_color(DARK))
            y += 24
            line = word
        else:
            line = next_line
    if line:
        draw.text((x1 + 16, y), line, font=font(18), fill=image_color(DARK))


def arrow(draw, start, end, label: str = ""):
    draw.line([start, end], fill=image_color(BLUE), width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        head = [(x2, y2), (x2 - 12 if x2 > x1 else x2 + 12, y2 - 8), (x2 - 12 if x2 > x1 else x2 + 12, y2 + 8)]
    else:
        head = [(x2, y2), (x2 - 8, y2 - 12 if y2 > y1 else y2 + 12), (x2 + 8, y2 - 12 if y2 > y1 else y2 + 12)]
    draw.polygon(head, fill=image_color(BLUE))
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        draw.text((mx - 35, my - 22), label, font=font(15, True), fill=image_color(BLUE))


def build_architecture_diagram() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1600, 1040), "white")
    draw = ImageDraw.Draw(image)
    draw.text((48, 28), "NXT desktop-first logical architecture", font=font(32, True), fill=image_color(NAVY))
    draw.text((48, 73), "Windows local runtime | Binance USD-M is the execution system of record", font=font(18), fill=image_color(MID))

    draw_box(draw, (60, 170, 360, 350), "Trader / Owner", "Local app review Approve / Reject Emergency controls", PALE_AMBER)
    draw_box(draw, (60, 510, 360, 700), "Telegram", "Signal alerts Approval status Critical notifications", PALE_AMBER)
    draw_box(draw, (550, 145, 1030, 350), "Desktop Application Layer", "Local UI API Approval workflow Lifecycle state Access control", PALE_BLUE)
    draw_box(draw, (550, 460, 1030, 760), "NXT Core Services", "Scheduler + scan orchestrator Market data adapter Rule engine Risk & proposal engine Position monitor Reconciliation", PALE_GREEN)
    draw_box(draw, (550, 845, 1030, 985), "Durable Local Store", "SQLite WAL: business entities, append-only audit events, checkpoints and idempotency keys", LIGHT_GRAY)
    draw_box(draw, (1200, 140, 1540, 365), "Binance USD-M", "Public market data Exchange info Account / orders User data stream", PALE_BLUE)
    draw_box(draw, (1200, 525, 1540, 735), "Execution Gateway", "Signed REST Order lifecycle Protective orders Exchange reconciliation", PALE_GREEN)
    draw_box(draw, (1200, 845, 1540, 985), "Windows Security", "DPAPI / Credential Manager Least-privilege keys Local logs and backup", LIGHT_GRAY)

    arrow(draw, (360, 255), (550, 255), "review")
    arrow(draw, (360, 605), (550, 605), "notify")
    arrow(draw, (790, 350), (790, 460), "commands")
    arrow(draw, (790, 760), (790, 845), "events")
    arrow(draw, (1030, 250), (1200, 250), "read")
    arrow(draw, (1030, 610), (1200, 610), "orders")
    arrow(draw, (1370, 365), (1370, 525), "fills")
    arrow(draw, (1030, 915), (1200, 915), "secrets")
    image.save(ARCHITECTURE_PNG)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_run(run, size=10.5, bold=False, color=DARK, italic=False, font_name="Aptos") -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], fill=LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    prevent_row_split(header)
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(text), size=9.25, bold=True)
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for i, text in enumerate(values):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(text), size=9.25, bold=False)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_para(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        style_run(p.add_run(bold_prefix), bold=True)
        style_run(p.add_run(text[len(bold_prefix):]))
    else:
        style_run(p.add_run(text))
    return p


def add_bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        add_para(doc, value, style="List Bullet")


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(title), bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    style_run(p2.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 6),
        ("Subtitle", 13, MID, 0, 12),
        ("Heading 1", 16, NAVY, 14, 7),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 11, TEAL, 8, 4),
    ]:
        s = styles[name]
        s.font.name = "Aptos Display" if name == "Title" else "Aptos"
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet",):
        s = styles[name]
        s.font.name = "Aptos"
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.42)
        s.paragraph_format.first_line_indent = Inches(-0.2)
        s.paragraph_format.space_after = Pt(3)
    header = section.header.paragraphs[0]
    style_run(header.add_run("NXT Automated Trading System | High Level Design"), size=8.5, bold=True, color=MID)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(footer.add_run("Desktop-first | HLD v0.1"), size=8.5, color=MID)


def build() -> None:
    build_architecture_diagram()
    doc = Document()
    configure(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(title.add_run("HIGH LEVEL DESIGN"), size=24, bold=True, color=NAVY, font_name="Aptos Display")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(subtitle.add_run("NXT Automated Signal and Trading System"), size=13, color=MID)
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ["Phiên bản", "v0.1 — Desktop-first architecture"],
        ["Ngày", date.today().strftime("%d/%m/%Y")],
        ["Nguồn yêu cầu", "BRD NXT Automated Signal and Trading System v0.2"],
        ["Strategy baseline", "NXT v3.5 USD-M BTCUSDT/BNBUSDT/SOLUSDT 1D; promoted rulebook 2026-07-14"],
        ["Runtime MVP", "Windows desktop: local app + Windows Task Scheduler + local durable store"],
        ["Execution boundary", "Binance USD-M Futures testnet first; live only after explicit approval and go-live controls"],
    ], [2200, 7160], PALE_BLUE)
    add_callout(doc, "Architecture decision", "MVP is a desktop-first modular monolith. Logical modules are independently testable, but are deployed locally together to minimize operational complexity. Cloud hosting is deliberately out of MVP scope and remains a later readiness decision.", PALE_AMBER)

    doc.add_heading("1. Purpose and scope", level=1)
    add_para(doc, "This HLD defines the target architecture for a human-in-the-loop NXT system: it scans a promoted NXT ruleset, produces an approved order proposal, executes only within policy, monitors real exchange state, and records an auditable lifecycle.")
    add_table(doc, ["In scope", "Out of scope / deferred"], [[
        "Desktop scan, rule evaluation, approval workflow, USD-M testnet gateway, order protection, reconciliation, audit, Telegram/local UI.",
        "Cloud runtime, autonomous rule promotion, HFT/intraday, multi-exchange routing, withdrawal automation, unattended live entry approval.",
    ]], [4680, 4680], PALE_BLUE)

    doc.add_heading("2. Architecture drivers and principles", level=1)
    add_table(doc, ["Driver", "Architectural response"], [
        ["Human approval", "Proposal and Decision are durable records. Execution accepts only a non-expired approved proposal hash."],
        ["NXT latest changes", "RuleVersion is immutable with parameters, artifact hash, source contract/session, and promotion timestamp. Existing positions retain their entry rule version."],
        ["USD-M execution risk", "Risk engine sizes before leverage; isolated margin, exchange filters, liquidation buffer and portfolio caps are pre-trade gates."],
        ["Ambiguous exchange outcomes", "Stable clientOrderId, query-before-retry, idempotency ledger and Binance reconciliation prevent duplicate logical orders."],
        ["Desktop-first reliability", "Task Scheduler invokes a controlled service command; local health status, StartWhenAvailable, checkpoint recovery and operator alerting cover host interruptions."],
        ["Auditability", "Append-only AuditEvent records actor, source, before/after state, correlation ID and external order IDs."],
    ], [2600, 6760])

    doc.add_heading("3. Logical architecture", level=1)
    doc.add_picture(str(ARCHITECTURE_PNG), width=Inches(6.85))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(caption.add_run("Figure 1. Desktop-first logical architecture"), size=8.5, italic=True, color=MID)
    add_table(doc, ["Layer", "Components", "Responsibility"], [
        ["Experience", "Local App; Telegram notifier", "Review proposals, display lifecycle/audit, capture Approve/Reject and provide alerts. Telegram is notification-first; privileged decisions remain policy-controlled."],
        ["Application", "Workflow service; lifecycle manager; policy evaluator", "Owns state transitions, approval expiry, idempotency and exception-to-Needs Review behavior."],
        ["Strategy", "Market data adapter; NXT rule engine; RuleVersion registry", "Reads closed USD-M candles, computes indicators, evaluates E/P/G rules, stores pass/fail rationale and signal snapshot."],
        ["Trading", "Risk/proposal engine; execution gateway; position monitor", "Calculates quantity/margin, validates filters, submits approved orders, maintains protective orders and reconciles exchange events."],
        ["Data/control", "SQLite WAL store; audit log; settings/secrets; backup exporter", "Durable local source for business state; Binance remains source of truth for exchange order/fill/position facts."],
        ["External", "Binance USD-M APIs/streams; Telegram API", "Market data, exchange status, trading/account events, delivery of notifications."],
    ], [1550, 2750, 5060])

    doc.add_heading("4. Desktop deployment topology", level=1)
    add_table(doc, ["Node", "Processes / stores", "Operational responsibility"], [
        ["Windows desktop host", "NXT local app/service; scheduled scan command; SQLite database; encrypted secrets; local log and backup folder.", "Primary MVP runtime. Service must start after login/boot according to Task Scheduler configuration and publish health status."],
        ["Binance USD-M", "Public market data; exchange information; signed REST; user data stream.", "Execution system of record. All uncertain outcomes reconcile from Binance before another action."],
        ["Telegram", "Bot API outbound notifications and optional review deep-link.", "Secondary interaction channel. Notification failure must not lose persisted signal/proposal state."],
    ], [1900, 3900, 3560])
    add_callout(doc, "Desktop boundary", "No inbound public internet service is required for MVP. The local app binds to loopback by default. Remote access, public endpoints and cloud workers are not part of this HLD's deployment baseline.", PALE_GREEN)

    doc.add_heading("5. Core workflows", level=1)
    doc.add_heading("5.1 Daily scan and proposal", level=2)
    add_table(doc, ["Step", "Flow and control"], [
        ["1. Schedule", "Windows Task Scheduler starts the scan after USD-M 1D candle close (default 07:10 ICT). A ScanRun receives a correlation ID."],
        ["2. Data quality", "Market Data Adapter verifies symbol, source, closed candle, freshness and expected session. Stale/incomplete data produces no new signal and raises an operator alert."],
        ["3. Rule evaluation", "NXT Rule Engine loads the promoted immutable RuleVersion and creates a Signal snapshot only after all E/P/G conditions pass, including G-02."],
        ["4. Proposal", "Risk Engine reads account snapshot, open-risk cap, exchange filters and policy; it computes quantity, notional, isolated margin requirement, leverage constraint, SL/TP1 and expiry."],
        ["5. Persist then notify", "Workflow persists Signal and OrderProposal before Telegram/local UI notification. Signal deduplication uses versioned deterministic keys."],
        ["6. Decision", "Trader approves/rejects the exact proposal hash. A changed price, expired proposal, changed policy or changed levels invalidates approval."],
    ], [1000, 8360])
    doc.add_heading("5.2 Execution and protection", level=2)
    add_table(doc, ["Step", "Flow and control"], [
        ["1. Pre-trade revalidation", "Execution Gateway rechecks API permission, server time, account/margin, position state, leverage/margin mode, price deviation, filters, risk limits and idempotency."],
        ["2. Entry", "Gateway writes clientOrderId before submit. A timeout or 5xx is an unknown outcome: query order and user stream before retrying."],
        ["3. Fill handling", "Actual fills update Position and ExchangeOrder. Partial fill quantity becomes the maximum size for protective orders."],
        ["4. Protection", "Create reduce-only SL and TP1 immediately. Rejection creates Critical state, blocks new entries and follows owner-approved fail-safe policy."],
        ["5. Exit lifecycle", "Monitor detects TP1, Early-BE, stop and SSL runner action. Action is a recommendation or automated only at the approved automation level."],
    ], [1500, 7860])
    doc.add_heading("5.3 Monitoring and reconciliation", level=2)
    add_bullets(doc, [
        "User data stream is the low-latency feed for orders, fills and account/position changes; REST reconciliation repairs gaps after disconnect or restart.",
        "Binance is the source of truth for actual orders, fills and positions. Internal state is a reconciled projection, not a substitute for exchange facts.",
        "Every daily close runs a position recommendation using the entry RuleVersion or an explicit migration record. A new latest version cannot silently change an open position's policy.",
        "Any mismatch, ambiguous response, stale data or rejected protective order becomes Needs Review and suppresses dependent automation."],)

    doc.add_heading("6. Data architecture", level=1)
    add_table(doc, ["Entity", "Key contents", "Control"], [
        ["RuleVersion", "version, parameters, USD-M contract/session, system document hash, promotedAt", "Immutable and retained indefinitely."],
        ["ScanRun", "runId, candle date, symbols, source/freshness, status, errors, duration", "Unique run ID; operational retention at least two years."],
        ["Signal", "deterministic signalId, RuleVersion, indicator snapshot, guard rationale, levels", "Immutable once detected; deduplicated by source + rule version."],
        ["OrderProposal / Decision", "proposal hash, quantity, risk, expiry, decision actor/channel/time", "No in-place change after approval. New values require new proposal/decision."],
        ["ExchangeOrder / Position", "clientOrderId, Binance order ID, fills, protective order IDs, actual quantity/lifecycle", "Reconciled with exchange events and REST snapshots."],
        ["Recommendation / AuditEvent", "action rationale, automation level, status; actor, before/after, correlation ID", "Append-only; supports investigation and weekly audit."],
    ], [1800, 4750, 2810])
    add_callout(doc, "Local storage selection", "Proposed MVP store: SQLite in WAL mode with transactional writes, periodic encrypted backup/export and a migration layer. This meets the single-host desktop scope while preserving a future move to a managed database if cloud is approved.", PALE_AMBER)

    doc.add_heading("7. Interfaces and integration boundaries", level=1)
    add_table(doc, ["Interface", "Direction", "Purpose and resilience"], [
        ["USD-M market data", "Outbound read", "Fetch candles and exchange info. Validate closed bar and freshness; cache may support display only, never new signal generation when stale."],
        ["USD-M trading/account", "Outbound signed REST", "Query account/order status; submit/cancel/replace approved orders; use clientOrderId and query-before-retry."],
        ["USD-M user data stream", "Inbound event", "Receive account/order/fill events; reconnect and backfill via REST when sequence or connectivity is uncertain."],
        ["Telegram Bot API", "Outbound notification", "Signal, decision, fill and critical alerts. Retry delivery independently of business state persistence."],
        ["Local app API", "Loopback only", "UI reads business state and records permitted user actions. No public listener in MVP."],
    ], [2100, 1700, 5560])

    doc.add_heading("8. Security, risk and resilience", level=1)
    add_table(doc, ["Concern", "Design control"], [
        ["Secrets", "Store API credentials in Windows Credential Manager/DPAPI-backed storage; never in source, SQLite, exports or logs. Keys must lack withdrawal permission."],
        ["Environment separation", "Testnet and live use distinct configuration, credentials, clientOrderId namespace, database profile and visible UI banner."],
        ["Leverage and liquidation", "Risk budget determines quantity first. Leverage is a margin-efficiency setting, capped by policy; pre-trade validation verifies isolated margin and liquidation buffer before approval/execution."],
        ["Failure handling", "Unknown exchange outcome, stale data, protective failure or reconciliation mismatch transitions to Needs Review/Critical rather than inferring success."],
        ["Kill switch", "Pause blocks new entry while monitoring/protection continues. Kill switch blocks automated action except an owner-approved emergency policy; all changes are audited."],
        ["Desktop continuity", "Task Scheduler StartWhenAvailable, health check, alerting, checkpointed ScanRun and reconciliation on restart. Cloud is not the fallback in MVP."],
    ], [2600, 6760])

    doc.add_heading("9. NFR and requirement traceability", level=1)
    add_table(doc, ["BRD area", "HLD response", "Verification evidence"], [
        ["FR-01 / FR-02; BR-01 to BR-11", "Scheduler, data-quality gate, NXT Rule Engine and RuleVersion registry.", "Golden-signal test; G-02 regression; deterministic Signal ID."],
        ["FR-03 to FR-07; BR-12 to BR-15", "Risk/Proposal Engine, Approval workflow, Execution Gateway and Protective Order manager.", "Testnet E2E; pre-trade checks; partial fill and protective failure tests."],
        ["FR-08 / FR-09; BR-16 / BR-17", "Position monitor, user data adapter, REST reconciliation and recommendation service.", "Disconnect/restart/manual intervention reconciliation evidence."],
        ["FR-10 to FR-12; NFR-02/03/05/07", "Notifier, local dashboard, audit store, health monitor and operational controls.", "Delivery retry logs, audit timeline, kill-switch UAT and scheduled-run status."],
        ["Desktop-first decision", "Single-host modular deployment and local data store; cloud excluded from MVP boundary.", "30-day desktop paper/testnet operational run."],
    ], [2250, 4570, 2540])

    doc.add_heading("10. Key design decisions and open decisions", level=1)
    add_table(doc, ["Decision", "Status", "Rationale / remaining action"], [
        ["Desktop-first modular monolith", "Confirmed", "Lowest operational complexity for MVP; components remain separable for later cloud assessment."],
        ["Binance USD-M as technical baseline", "Confirmed", "Matches promoted NXT data and funding model. Owner must still confirm account eligibility/contract availability before testnet/live."],
        ["Human approval for entry", "Confirmed", "No live entry without an approval record matching the proposal hash and environment."],
        ["Live risk, leverage ceiling and loss caps", "Open / blocking", "Owner must set per-trade and portfolio policies. HLD enforces a policy interface rather than hardcoding values."],
        ["Runner exit automation level", "Open / high", "Recommendation-only is the safe default until owner approves an automation policy."],
        ["Cloud readiness", "Deferred", "Evaluate only after desktop paper/testnet pilot achieves the BRD operational-run exit criteria."],
    ], [2900, 1750, 4710])

    doc.add_heading("11. Phased implementation view", level=1)
    add_table(doc, ["Phase", "Architecture increment", "Exit condition"], [
        ["1. Desktop paper", "SQLite schema, RuleVersion registry, scan workflow, local UI/Telegram approval, simulated exchange adapter and audit timeline.", "Golden signal, approval, duplicate and restart UAT pass."],
        ["2. Desktop USD-M testnet", "Signed gateway, exchange filters, user stream, protective orders, reconciliation, risk/portfolio pre-trade gates.", "30-day stable run without duplicate order or unresolved critical protection failure."],
        ["3. Desktop live assisted", "Feature-gated live profile, manual entry approval, production monitoring, backup/restore drill and low-risk policy configuration.", "Go-live checklist and owner approval."],
        ["4. Later assessment", "Optional managed automation and cloud-readiness design only after desktop evidence supports it.", "Separate decision record; no automatic migration."],
    ], [1750, 5180, 2430])

    doc.add_heading("12. HLD acceptance criteria", level=1)
    add_bullets(doc, [
        "Architecture supports every Must requirement in BRD v0.2 without requiring cloud deployment.",
        "Signals, proposals, decisions and exchange actions are traceable by correlation ID and immutable rule/proposal references.",
        "No execution path bypasses approval, pre-trade checks, idempotency or protective-order validation.",
        "Rule update, G-02 regression, stale-data block, restart recovery and manual Binance action have explicit handling paths.",
        "All open risk, leverage, margin and automation values remain policy-configured and require owner approval before live use."],)

    doc.add_heading("Appendix A. References", level=1)
    add_bullets(doc, [
        "docs/BRD_NXT_Automated_Signal_Trading_System_v0.2.docx",
        "latest/NXT_Latest_Summary.md",
        "latest/NXT_Latest_NXT35_USDM_BlockShortAfterLosingLong_System_And_Indicators.docx",
        "latest/NXT_Latest_NXT35_USDM_BlockShortAfterLosingLong_SignalRegression.json",
        "app/nxt_signal_app.py and scripts/daily_nxt_signal_scan.py — current desktop scan core.",
    ])

    props = doc.core_properties
    props.title = "HLD - NXT Automated Signal and Trading System"
    props.subject = "Desktop-first high-level architecture based on BRD v0.2"
    props.author = "System Architecture"
    props.comments = "HLD v0.1; not authorization for live trading."
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
