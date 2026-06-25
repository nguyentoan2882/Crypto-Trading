from __future__ import annotations

import json
import shutil
import sys
import time
import urllib.request
from datetime import date, datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

sys.path.insert(0, str(Path(__file__).resolve().parent))

import backtest_nxt31_utc7_latest as base
import backtest_nxt32_native_1d_latest as native
import test_nxt33_long_only_pullback_continuation as cont
from test_nxt33_ssl14 import enrich_with_ssl_period


ROOT = Path(__file__).resolve().parents[1]
CACHE = ROOT / "data_cache" / "binance_spot_1d"
OUT_DIR = ROOT / "outputs" / "nxt35_btc_bnb_sol_latest"
OUT_JSON = OUT_DIR / "nxt35_btc_bnb_sol_latest_results.json"
OUT_XLSX = OUT_DIR / "NXT35_BTC_BNB_SOL_20K.xlsx"
LATEST = ROOT / "latest"
ARCHIVE = ROOT / "outputs" / "archive_from_latest"
LATEST_JSON = LATEST / "NXT_Latest_NXT35_BTC_BNB_SOL_LongOnlyPullbackContinuation_20K.json"
LATEST_XLSX = LATEST / "NXT_Latest_NXT35_BTC_BNB_SOL_LongOnlyPullbackContinuation_20K.xlsx"
LATEST_DOCX = LATEST / "NXT_Latest_NXT35_BTC_BNB_SOL_System_And_Indicators.docx"
LATEST_SUMMARY = LATEST / "NXT_Latest_Summary.md"

SYMBOLS = ["BTCUSDT", "BNBUSDT", "SOLUSDT"]
START_DATE = native.START_DATE
END_DATE = native.END_DATE
WARMUP_DATE = native.WARMUP_DATE
STARTING_EQUITY = 20_000
ONE_R_DOLLARS = 1_000
SYSTEM_VERSION = "NXT v3.5 Portfolio BTC+BNB+SOL + Binance Native 1D + SSL14 + Runner A + Anti-Immediate-Reversal + LONG-only Pullback Continuation + No Risk-Off"

RULES = [
    "Data: Binance native 1D candles for BTCUSDT, BNBUSDT, and SOLUSDT.",
    "ATR14 uses the NXT v3.5 ATR-SMA variant.",
    "SSL Channel: SMA(high,14) and SMA(low,14); state flips bullish when close is above high SMA and bearish when close is below low SMA.",
    "Primary LONG: SSL flips bullish, price crosses above EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 > 50.",
    "Primary SHORT: SSL flips bearish, price crosses below EMA20 within the last 3 candles, distance from close to EMA50 <= 2 ATR14, and RSI14 < 50.",
    "Continuation LONG: SSL is bullish, close > EMA20 > EMA50, low touched EMA20 within the last 5 candles, close > EMA20, and close > previous close.",
    "Continuation is LONG-only; SHORT continuation is disabled.",
    "Continuation does not require RSI, distance-to-EMA50, or EMA50 slope filters.",
    "Anti-immediate-reversal: after a profitable runner exits by opposite SSL flip, block an opposite-direction entry on the exit candle and the next candle.",
    "Initial stop: 1.5 ATR14 from entry.",
    "TP1: 2.5 ATR14 from entry; close 50% at TP1.",
    "Runner A: after TP1, move remaining 50% stop to breakeven and exit runner on opposite SSL flip or breakeven stop.",
    "Risk-off overlay is disabled.",
    "Round-trip trading cost is included in R results.",
    "20K account model: starting equity $20,000 and 1R = $1,000.",
]


def date_label(ms: int) -> str:
    return datetime.fromtimestamp(ms / 1000, timezone.utc).date().isoformat()


def fetch_binance_native_1d(symbol: str) -> list[dict]:
    CACHE.mkdir(parents=True, exist_ok=True)
    path = CACHE / f"{symbol}.json"
    if path.exists():
        rows = json.loads(path.read_text(encoding="utf-8"))
    else:
        rows = []
    if rows:
        last = max(int(r["time"]) for r in rows)
    else:
        last = int(datetime(WARMUP_DATE.year, WARMUP_DATE.month, WARMUP_DATE.day, tzinfo=timezone.utc).timestamp() * 1000) - 86_400_000

    end_ms = int(datetime(END_DATE.year, END_DATE.month, END_DATE.day, tzinfo=timezone.utc).timestamp() * 1000)
    start = max(last + 86_400_000, int(datetime(WARMUP_DATE.year, WARMUP_DATE.month, WARMUP_DATE.day, tzinfo=timezone.utc).timestamp() * 1000))
    while start <= end_ms:
        url = f"https://data-api.binance.vision/api/v3/klines?symbol={symbol}&interval=1d&startTime={start}&endTime={end_ms}&limit=1000"
        with urllib.request.urlopen(url, timeout=30) as response:
            batch = json.loads(response.read().decode("utf-8"))
        if not batch:
            break
        for item in batch:
            rows.append({
                "time": int(item[0]),
                "open": float(item[1]),
                "high": float(item[2]),
                "low": float(item[3]),
                "close": float(item[4]),
                "volume": float(item[5]),
            })
        next_start = int(batch[-1][0]) + 86_400_000
        if next_start <= start:
            break
        start = next_start
        time.sleep(0.03)

    rows = sorted({int(r["time"]): r for r in rows}.values(), key=lambda r: int(r["time"]))
    path.write_text(json.dumps(rows), encoding="utf-8")
    out = []
    for row in rows:
        d = date_label(int(row["time"]))
        if WARMUP_DATE <= date.fromisoformat(d) <= END_DATE:
            item = dict(row)
            item["localDate"] = d
            out.append(item)
    return out


def archive_existing_latest() -> None:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    target = ARCHIVE / f"before_btc_bnb_sol_latest_{stamp}"
    target.mkdir(parents=True, exist_ok=True)
    for path in LATEST.glob("NXT_Latest_*"):
        if path.is_file() and path.name != "NXT_Latest_Summary.md":
            shutil.move(str(path), str(target / path.name))


def equity_curve(trades: list[dict]) -> list[dict]:
    equity = STARTING_EQUITY
    peak = equity
    rows = []
    for i, trade in enumerate(trades, 1):
        pnl = trade["rMultiple"] * ONE_R_DOLLARS
        equity += pnl
        peak = max(peak, equity)
        rows.append({
            "no": i,
            "exitTime": trade["exitTime"],
            "symbol": trade["symbol"],
            "side": trade["side"],
            "signalType": trade["signalType"],
            "rMultiple": trade["rMultiple"],
            "pnl": pnl,
            "equity": equity,
            "drawdown": equity - peak,
        })
    return rows


def write_row(ws, r: int, values: list) -> None:
    for c, value in enumerate(values, 1):
        ws.cell(r, c).value = value


def style_sheet(ws, header_row: int = 4) -> None:
    ws.freeze_panes = f"A{header_row + 1}"
    for cell in ws[header_row]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = min(max(ws.column_dimensions[get_column_letter(col)].width or 12, 12), 26)


def build_workbook(result: dict) -> None:
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    summary["A1"] = "NXT v3.5 Latest Portfolio - BTC BNB SOL"
    summary["A2"] = "Starting account $20,000 | 1R = $1,000 | Binance native 1D | ATR-SMA | LONG-only pullback continuation."
    write_row(summary, 4, ["Metric", "Value"])
    stats = result["stats"]
    rows = [
        ("Symbols", ", ".join(result["symbols"])),
        ("Trades", stats["trades"]),
        ("Wins", stats["wins"]),
        ("Losses", stats["losses"]),
        ("Win Rate", stats["winRate"]),
        ("Total R", stats["totalR"]),
        ("Average R", stats["avgR"]),
        ("Max DD R", stats["maxDrawdownR"]),
        ("Best R", stats["bestR"]),
        ("Worst R", stats["worstR"]),
        ("Profit Factor", stats["profitFactor"]),
        ("Starting Equity", STARTING_EQUITY),
        ("1R Dollars", ONE_R_DOLLARS),
        ("Ending Equity", result["account"]["endingEquity"]),
        ("Net Profit", result["account"]["netProfit"]),
        ("Max DD Dollars", result["account"]["maxDrawdownDollars"]),
    ]
    for r, row in enumerate(rows, 5):
        write_row(summary, r, list(row))
    style_sheet(summary)
    summary.column_dimensions["A"].width = 24
    summary.column_dimensions["B"].width = 34

    trades_ws = wb.create_sheet("Trades")
    headers = ["Symbol", "No", "Signal Type", "Side", "Signal Date", "Entry Date", "Entry Price", "Initial Stop", "Final Stop", "Risk / Unit", "TP1", "TP1 Date", "Exit Date", "Exit Price", "Exit Reason", "R", "P&L $", "Equity $", "ATR14", "RSI14", "Distance EMA50 ATR", "Notes"]
    write_row(trades_ws, 4, headers)
    curve = result["equityCurve"]
    equity_by_key = {(r["symbol"], r["exitTime"], r["no"]): r for r in curve}
    for i, trade in enumerate(result["trades"], 5):
        eq = curve[i - 5]
        write_row(trades_ws, i, [
            trade["symbol"].replace("USDT", ""), trade["tradeNo"], trade["signalType"], trade["side"],
            trade["signalTime"], trade["entryTime"], trade["entryPrice"], trade["initialStop"], trade["finalStop"],
            trade["riskPerUnit"], trade["tp1"], trade["tp1Time"], trade["exitTime"], trade["exitPrice"],
            trade["exitReason"], trade["rMultiple"], eq["pnl"], eq["equity"], trade["atr14"], trade["rsi14"],
            trade["distanceToEma50Atr"], trade["notes"],
        ])
    style_sheet(trades_ws)
    trades_ws.column_dimensions["V"].width = 42

    eq_ws = wb.create_sheet("20K Account")
    eq_ws["A1"] = "20K Account Equity Curve"
    write_row(eq_ws, 4, ["No", "Exit Date", "Symbol", "Side", "Signal Type", "R", "P&L $", "Equity $", "Drawdown $"])
    for r, row in enumerate(curve, 5):
        write_row(eq_ws, r, [row["no"], row["exitTime"], row["symbol"].replace("USDT", ""), row["side"], row["signalType"], row["rMultiple"], row["pnl"], row["equity"], row["drawdown"]])
    style_sheet(eq_ws)

    by_symbol = wb.create_sheet("By Symbol")
    write_row(by_symbol, 4, ["Symbol", "Trades", "Win Rate", "Total R", "Avg R", "Max DD R", "Best R", "Worst R", "Profit Factor"])
    for r, row in enumerate(result["bySymbol"], 5):
        write_row(by_symbol, r, [row["group"].replace("USDT", ""), row["trades"], row["winRate"], row["totalR"], row["avgR"], row["maxDrawdownR"], row["bestR"], row["worstR"], row["profitFactor"]])
    style_sheet(by_symbol)

    cont_ws = wb.create_sheet("Continuation")
    write_row(cont_ws, 4, ["Symbol", "Trades", "Win Rate", "Total R", "Avg R", "Max DD R", "Best R", "Worst R", "Profit Factor"])
    for r, row in enumerate(result["byContinuationSymbol"], 5):
        write_row(cont_ws, r, [row["group"].replace("USDT", ""), row["trades"], row["winRate"], row["totalR"], row["avgR"], row["maxDrawdownR"], row["bestR"], row["worstR"], row["profitFactor"]])
    style_sheet(cont_ws)

    assumptions = wb.create_sheet("Assumptions")
    write_row(assumptions, 4, ["#", "Assumption"])
    for r, line in enumerate(result["assumptions"], 5):
        write_row(assumptions, r, [r - 4, line])
    style_sheet(assumptions)
    assumptions.column_dimensions["B"].width = 78

    quality = wb.create_sheet("Data Quality")
    write_row(quality, 4, ["Symbol", "Daily Rows", "First Day", "Last Day", "Source"])
    for r, (sym, q) in enumerate(result["datasets"].items(), 5):
        write_row(quality, r, [sym.replace("USDT", ""), q["dailyRows"], q["firstDay"], q["lastDay"], q["source"]])
    style_sheet(quality)

    wb.save(OUT_XLSX)


def build_docx(result: dict) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("NXT v3.5 Latest Portfolio - BTC BNB SOL", 0)
    doc.add_paragraph("Current selected latest portfolio: BTCUSDT, BNBUSDT, and SOLUSDT.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    for key, value in [
        ("System", result["systemVersion"]),
        ("Symbols", ", ".join(result["symbols"])),
        ("Trades", result["stats"]["trades"]),
        ("Total R", f"{result['stats']['totalR']:.2f}R"),
        ("Max DD R", f"{result['stats']['maxDrawdownR']:.2f}R"),
        ("Win rate", f"{result['stats']['winRate']:.2%}"),
        ("Profit factor", f"{result['stats']['profitFactor']:.2f}"),
        ("Starting equity", f"${STARTING_EQUITY:,.2f}"),
        ("Ending equity", f"${result['account']['endingEquity']:,.2f}"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    doc.add_heading("Rules", level=1)
    for line in RULES:
        doc.add_paragraph(line, style="List Bullet")
    doc.save(LATEST_DOCX)


def build_result() -> dict:
    all_trades = []
    datasets = {}
    for symbol in SYMBOLS:
        candles = enrich_with_ssl_period(fetch_binance_native_1d(symbol), 14)
        datasets[symbol] = {
            "dailyRows": len(candles),
            "firstDay": candles[0]["localDate"],
            "lastDay": candles[-1]["localDate"],
            "source": "Binance spot native 1D klines",
        }
        all_trades.extend(cont.backtest_symbol(symbol, candles))
    all_trades.sort(key=lambda trade: trade["exitTime"])
    stats = cont.enriched_stats(all_trades)
    curve = equity_curve(all_trades)
    max_dd_dollars = min((row["drawdown"] for row in curve), default=0)
    result = {
        "generatedAt": datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "systemVersion": SYSTEM_VERSION,
        "period": {
            "start": START_DATE.isoformat(),
            "end": (END_DATE - base.timedelta(days=1)).isoformat(),
            "timezone": "Binance native daily candles",
        },
        "symbols": SYMBOLS,
        "stats": stats,
        "account": {
            "startingEquity": STARTING_EQUITY,
            "oneRDollars": ONE_R_DOLLARS,
            "endingEquity": STARTING_EQUITY + stats["totalR"] * ONE_R_DOLLARS,
            "netProfit": stats["totalR"] * ONE_R_DOLLARS,
            "maxDrawdownDollars": max_dd_dollars,
        },
        "trades": all_trades,
        "continuationStats": cont.enriched_stats([t for t in all_trades if t["signalType"] == "Continuation"]),
        "bySymbol": cont.group_stats(all_trades, lambda trade: trade["symbol"]),
        "byContinuationSymbol": cont.group_stats([t for t in all_trades if t["signalType"] == "Continuation"], lambda trade: trade["symbol"]),
        "equityCurve": curve,
        "datasets": datasets,
        "assumptions": RULES,
    }
    return result


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    LATEST.mkdir(parents=True, exist_ok=True)
    result = build_result()
    OUT_JSON.write_text(json.dumps(result, indent=2), encoding="utf-8")
    build_workbook(result)
    archive_existing_latest()
    shutil.copy2(OUT_JSON, LATEST_JSON)
    shutil.copy2(OUT_XLSX, LATEST_XLSX)
    build_docx(result)
    LATEST_SUMMARY.write_text(
        "\n".join([
            "# Latest NXT System",
            "",
            f"System: {result['systemVersion']}",
            "",
            f"Symbols: {', '.join(result['symbols'])}",
            f"Trades: {result['stats']['trades']}",
            f"Total R: {result['stats']['totalR']:.2f}R",
            f"Max DD R: {result['stats']['maxDrawdownR']:.2f}R",
            f"Win rate: {result['stats']['winRate']:.2%}",
            f"Profit factor: {result['stats']['profitFactor']:.2f}",
            f"Starting equity: ${STARTING_EQUITY:,.2f}",
            f"1R: ${ONE_R_DOLLARS:,.2f}",
            f"Ending equity: ${result['account']['endingEquity']:,.2f}",
            f"Net profit: ${result['account']['netProfit']:,.2f}",
            f"Max DD dollars: ${result['account']['maxDrawdownDollars']:,.2f}",
            "",
            f"Continuation trades: {result['continuationStats']['trades']}",
            f"Continuation R: {result['continuationStats']['totalR']:.2f}R",
            "",
            "Notes: Latest selected portfolio is BTCUSDT, BNBUSDT, and SOLUSDT. Uses NXT v3.5 ATR-SMA logic, profitable-runner Anti-Immediate-Reversal, and LONG-only pullback/touch EMA20 continuation.",
            "",
            f"Workbook: {LATEST_XLSX.name}",
            f"JSON: {LATEST_JSON.name}",
            f"System doc: {LATEST_DOCX.name}",
        ]),
        encoding="utf-8",
    )
    print(json.dumps({"latestJson": str(LATEST_JSON), "latestXlsx": str(LATEST_XLSX), "stats": result["stats"], "account": result["account"]}, indent=2))


if __name__ == "__main__":
    main()
